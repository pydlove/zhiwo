#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a DOCX article in 清新律动风 style.
"""
import os
import re
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT

# ── paths ──────────────────────────────────────────────────────────
IMG_DIR = "/tmp/article_a1_imgs"
OUT_DIR = "/Users/panyong/aio_project/小程序/docs/文章/2026-05-01/wangcl"
OUT_FILE = os.path.join(OUT_DIR, "45岁二婚朋友的忠告：别等离婚了才懂这3个道理.docx")

# ── style colours ──────────────────────────────────────────────────
TEAL      = RGBColor(0x2A, 0x9D, 0x8F)
CORAL     = RGBColor(0xE7, 0x6F, 0x51)
DARK      = RGBColor(0x26, 0x46, 0x53)
GOLD      = RGBColor(0xE9, 0xC4, 0x6A)
SAGE      = RGBColor(0xA8, 0xDA, 0xDC)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ── helpers ────────────────────────────────────────────────────────
def set_run_font(run, font_name="微软雅黑", size_pt=None, bold=False, italic=False, color=None):
    """Set font with eastAsia support for CJK."""
    font = run.font
    if size_pt:
        font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_empty_para(doc, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    return p


def add_english_label(doc, text="LIFE RHYTHM"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, "Georgia", 10, color=SAGE)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, "微软雅黑", 28, bold=True, color=DARK)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "微软雅黑", 18, color=TEAL)
    return p


def add_tagline(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, "微软雅黑", 13, italic=True, color=CORAL)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("~ ~ ~")
    set_run_font(run, "Georgia", 14, color=GOLD)
    return p


def add_chapter_marker(doc, num, title_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run1 = p.add_run(f"{num:02d}")
    set_run_font(run1, "Arial", 24, bold=True, color=TEAL)
    run2 = p.add_run(" / ")
    set_run_font(run2, "Arial", 16, color=RGBColor(0xCC, 0xCC, 0xCC))
    run3 = p.add_run(title_text)
    set_run_font(run3, "微软雅黑", 14, bold=True, color=DARK)
    return p


def add_body(doc, text, highlight_phrases=None):
    """Add body paragraph with optional teal-highlighted key phrases."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if not highlight_phrases:
        run = p.add_run(text)
        set_run_font(run, "微软雅黑", 12, color=DARK)
        return p

    # simple highlight by splitting
    remaining = text
    for phrase in highlight_phrases:
        if phrase in remaining:
            parts = remaining.split(phrase, 1)
            if parts[0]:
                r = p.add_run(parts[0])
                set_run_font(r, "微软雅黑", 12, color=DARK)
            r = p.add_run(phrase)
            set_run_font(r, "微软雅黑", 12, bold=True, color=TEAL)
            remaining = parts[1] if len(parts) > 1 else ""
    if remaining:
        r = p.add_run(remaining)
        set_run_font(r, "微软雅黑", 12, color=DARK)
    return p


def add_quote_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    # left border simulation via shading / paragraph border not directly exposed,
    # we use a light background and indent to mimic.
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, "微软雅黑", 13, italic=True, color=CORAL)
    # We can't easily set paragraph left border in python-docx without raw XML;
    # use background shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        from docx.oxml import parse_xml
        shd = parse_xml(r'<w:shd {} w:fill="FEF9F0" w:val="clear"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ))
        pPr.append(shd)
    else:
        shd.set(qn('w:fill'), "FEF9F0")
    return p


def add_highlight_box(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run1 = p.add_run("✦ ")
    set_run_font(run1, "微软雅黑", 12, bold=True, color=TEAL)
    run2 = p.add_run(text)
    set_run_font(run2, "微软雅黑", 12, bold=True, color=DARK)
    # background
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        from docx.oxml import parse_xml
        shd = parse_xml(r'<w:shd {} w:fill="F0FAF4" w:val="clear"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ))
        pPr.append(shd)
    else:
        shd.set(qn('w:fill'), "F0FAF4")
    return p


def add_image_centered(doc, img_path, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    # caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cr = cap.add_run(f"▲ {caption_text}")
    set_run_font(cr, "Georgia", 10, italic=True, color=SAGE)
    return p


def add_end_tags(doc, tags):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("  ".join([f"#{t}" for t in tags]))
    set_run_font(run, "微软雅黑", 11, color=TEAL)
    return p


# ── forbidden words map ────────────────────────────────────────────
FORBIDDEN = {
    "永久": "长期", "关注": "留意", "最新": "新近", "最": "很",
    "最大": "很大", "群": "圈子", "第一": "领先", "唯一": "少有的",
    "顶级": "优质", "最高级": "更高级", "万能": "多用途", "100%": "绝大多数",
    "绝对": "相当", "最好": "很好", "最佳": "很出色", "最强": "很强",
    "首选": "优选", "全网最低": "价格优惠", "最便宜": "性价比高",
    "国家级": "省级以上", "首家": "率先", "独家": "特有",
    "史上最": "非常", "特效": "效果明显", "震惊": "令人关注",
    "跳楼价": "超值价", "亏本卖": "薄利多销", "错过再等一年": "限时优惠",
    "保证治愈": "有望改善", "药到病除": "有助缓解", "治愈": "改善",
    "根治": "从根本上改善", "无副作用": "成分温和", "疗效": "效果",
    "秘方": "传统方法", "偏方": "传统方法", "神医": "资深医师",
    "神药": "有效药物", "无效退款": "可咨询售后", "抗癌": "辅助调理",
    "增高": "促进发育", "一吃就瘦": "配合运动", "减肥神药": "辅助产品",
    "延寿": "健康养生", "几天见效": "坚持服用", "暴富": "财富增长",
    "保本保息": "稳健型", "稳赚不赔": "收益相对稳定", "保本": "稳健型",
    "稳赚": "收益相对稳定", "零风险": "风险可控", "原始股": "股权投资",
    "内幕消息": "市场分析", "配资": "融资服务", "收益率100%": "预期收益",
    "高回报": "预期收益较好", "假一赔十": "正品保障", "免费送": "赠品活动",
    "烟草": "其他商品", "爆炸性": "重要", "突发": "刚刚", "重磅": "重要",
    "速看": "值得关注", "躺赚": "被动收入", "限时免费": "免费体验",
    "零元购": "免费试用", "全网首发": "率先发布", "白菜价": "性价比高",
    "先到先得": "数量有限", "转发此文章": "欢迎分享", "点击关注": "欢迎关注",
    "最后机会": "机会难得", "算命": "运势分析", "改运": "积极心态",
    "约炮": "交友", "一夜情": "短期关系", "赌博": "娱乐",
    "彩票预测": "彩票分析", "时时彩": "彩票", "外挂": "辅助工具",
    "走私": "跨境贸易", "破解版": "正版授权", "盗版": "非正版",
    "VPN": "网络工具", "翻墙": "跨境访问", "枪支": "器械",
    "代购": "海外购", "占卜": "性格测试", "八字": "出生日期",
    "风水": "环境布局", "开光": "仪式", "写真": "照片",
    "命格": "性格特点", "转运": "好运气", "荐股": "投资建议",
    "电子烟": "雾化器", "百分百": "绝大多数",
}
EXTRA_FORBIDDEN = ["分享到朋友圈", "不转不是中国人", "不转发死全家", "转发好运", "福利姬", "无码", "有码"]


def sanitize_text(text):
    for bad, good in FORBIDDEN.items():
        text = text.replace(bad, good)
    for bad in EXTRA_FORBIDDEN:
        text = text.replace(bad, "")
    return text


def count_chars_in_doc(doc):
    total = 0
    for para in doc.paragraphs:
        total += len(para.text)
    return total


# ── build document ─────────────────────────────────────────────────
def build_document():
    doc = Document()
    # narrow margins for a magazine feel
    sections = doc.sections[0]
    sections.top_margin = Cm(2)
    sections.bottom_margin = Cm(2)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    # ── header label ──
    add_english_label(doc, "LIFE RHYTHM")

    # ── title ──
    add_title(doc, "45岁二婚朋友的忠告：别等离婚了才懂这3个道理")

    # ── subtitle ──
    add_subtitle(doc, "婚姻不是终点，成长才是")

    # ── tagline ──
    add_tagline(doc, "她用十年弯路换来的清醒，希望你不必重蹈覆辙")

    add_empty_para(doc, Pt(6))

    # ── opening image ──
    add_image_centered(doc, os.path.join(IMG_DIR, "img1.jpg"),
                       "晨光中的咖啡与书，是独处也是沉淀")

    # ── opening body ──
    add_body(doc,
        "上周和老友林姐约在咖啡馆见面。她今年四十五岁，二婚三年，眉眼间却比从前松弛了很多。"
        "聊起婚姻，她放下杯子轻声说：\"如果头婚时有人告诉我这些，也许就不用绕那么大一圈了。\""
        "她的语气里没有怨恨，只有一种历经波澜后的平静。那天下午，阳光透过玻璃窗落在桌面上，"
        "她缓缓道来三个在两次婚姻里才慢慢读懂的道理。"
        "我坐在对面，看着她眼角细微的皱纹和嘴角淡淡的笑意，忽然意识到，"
        "真正的成熟不是 never 跌倒，而是跌倒后愿意爬起来，拍拍身上的灰尘，"
        "然后告诉自己：下一次，我要走得更稳一些。"
        "林姐就是这样的人。她的头婚持续了十二年，从校园爱情走到柴米油盐，"
        "最终却在日复一日的消磨中分道扬镳。那段日子对她来说像是一场漫长的雨季，"
        "潮湿、阴冷，看不到尽头。但正是那场雨，让她在后来的日子里学会了如何为自己撑伞，"
        "也学会了如何在晴天里好好珍惜阳光。",
        highlight_phrases=["历经波澜后的平静", "真正的成熟不是 never 跌倒，而是跌倒后愿意爬起来"]
    )

    add_divider(doc)

    # ── Chapter 1 ──
    add_chapter_marker(doc, 1, "别把伴侣当成情绪垃圾桶")

    add_body(doc,
        "林姐说，她头婚时有个很不好的习惯：工作受了委屈、和父母闹了矛盾、"
        "甚至只是地铁上被人踩了一脚，回家都要一股脑倒给丈夫。"
        "她以为这是亲密，是信任，是\"我把你当自己人才这样\"。"
        "可时间久了，丈夫从耐心倾听变成敷衍应付，最后干脆沉默以对。"
        "她当时很委屈，觉得对方\"变了\"，却没意识到，自己的情绪风暴早已把婚姻吹得千疮百孔。",
        highlight_phrases=["情绪风暴早已把婚姻吹得千疮百孔"]
    )

    add_quote_block(doc,
        "\"婚姻是两个成年人的结盟，不是一方对另一方的无限索取。\""
    )

    add_body(doc,
        "二婚之后，林姐学会了先处理情绪，再回家说话。"
        "难过的时候，她会先去公园走一圈，或者找个安静的角落把情绪写下来。"
        "等心里那团火熄了，再和伴侣沟通。"
        "她发现，当自己不再把婚姻当作情绪的泄洪口，对方反而更愿意靠近她、理解她。"
        "因为此时的交流不再是倾倒，而是分享。"
    )

    add_highlight_box(doc,
        "情绪稳定不是压抑自己，而是学会在开口之前，先给自己一个缓冲的空间。"
    )

    add_image_centered(doc, os.path.join(IMG_DIR, "img2.jpg"),
                       "安静的公园长椅，适合一个人整理心情")

    add_divider(doc)

    # ── Chapter 2 ──
    add_chapter_marker(doc, 2, "尊重比爱更难得")

    add_body(doc,
        "头婚时，林姐和丈夫是大学同学，恋爱长跑七年，感情基础看似很牢固。"
        "但结婚后她渐渐发现，他们之间的\"爱\"更像是一种惯性：习惯了彼此的存在，"
        "却忘了去尊重对方作为一个独立个体的需求和边界。"
        "她喜欢周末宅家看书，他觉得\"没意思\"，非要拉她出去社交；"
        "他想换城市发展事业，她觉得\"不顾家\"，百般阻挠。"
        "两个人都在以爱的名义，试图把对方改造成自己想要的样子。"
        "林姐回忆说，那时候他们几乎每周都要争吵，不是为了什么大事，"
        "而是因为谁都不愿意让步，谁都觉得对方\"应该\"理解自己。"
        "她会在争吵后冷战三天，等他来道歉；他则觉得她在无理取闹，越来越不愿意回家。"
        "两个人的心，就在这样一次次的小摩擦中，慢慢磨出了裂痕。",
        highlight_phrases=["以爱的名义，试图把对方改造成自己想要的样子", "慢慢磨出了裂痕"]
    )

    add_body(doc,
        "二婚的丈夫是个性格温和的人，话不多，但有一个习惯让林姐很感动："
        "每次她提出不同的意见，他不会急着反驳，而是认真听完，然后说\"我理解你的想法\"。"
        "这句话听起来普通，却让她感受到一种久违的尊重。"
        "她不再需要为了证明自己而争吵，也不再担心表达真实想法会被否定。"
        "婚姻里有了尊重，爱才能真正扎根。"
    )

    add_quote_block(doc,
        "\"爱一个人，是欣赏他本来的样子；尊重一个人，是允许他和自己不一样。\""
    )

    add_highlight_box(doc,
        "真正的亲密，不是两个人变成一个人，而是两个独立的人，选择并肩前行。"
    )

    add_image_centered(doc, os.path.join(IMG_DIR, "img3.jpg"),
                       "并肩而行的背影，是尊重也是陪伴")

    add_divider(doc)

    # ── Chapter 3 ──
    add_chapter_marker(doc, 3, "婚姻需要经营，而不是顺其自然")

    add_body(doc,
        "头婚时，林姐对婚姻有一种近乎天真的想法：只要两个人相爱，其他问题都会迎刃而解。"
        "所以她从不主动沟通矛盾，觉得\"时间长了就好了\"；"
        "她也不在意节日的仪式感，觉得\"老夫老妻不需要这些虚的\"。"
        "结果，矛盾像滚雪球一样越积越大，感情却在日复一日的忽视中慢慢冷却。"
        "等到发现的时候，彼此已经像住在同一个屋檐下的陌生人。",
        highlight_phrases=["像住在同一个屋檐下的陌生人"]
    )

    add_body(doc,
        "二婚之后，林姐彻底改变了这个观念。"
        "她开始主动和丈夫约定\"每周聊天时间\"，不谈家务、不谈孩子，只谈彼此这一周的心情和想法。"
        "她会在纪念日准备一顿简单的晚餐，不是为了形式，而是为了提醒彼此："
        "\"在这个忙碌的世界里，我依然把你放在心上。\""
        "她说，婚姻就像一座花园，你不浇水、不除草、不施肥，它当然不会自己开出花来。"
        "她还会定期和丈夫一起回顾过去一个月的相处，聊聊哪些时刻让自己感到被爱，"
        "哪些时刻又让自己觉得孤单。这种看似刻意的\"经营\"，反而让他们的感情越来越深厚。"
        "因为当两个人都愿意为这段关系花心思，婚姻就不再是负担，而是一种共同创造的乐趣。"
    )

    add_quote_block(doc,
        "\"顺其自然\"是婚姻里很危险的词，它往往意味着\"我不想再努力了\"。"
    )

    add_highlight_box(doc,
        "好的婚姻，不是找到了对的人就万事大吉，而是两个人都愿意为这段关系持续投入心力。"
    )

    add_image_centered(doc, os.path.join(IMG_DIR, "img4.jpg"),
                       "精心布置的餐桌，是经营也是告白")

    add_divider(doc)

    # ── closing ──
    add_body(doc,
        "临走时，林姐握着我的手说：\"我不是劝谁离婚，而是希望每一对夫妻都能在还来得及的时候，"
        "好好审视自己的婚姻。\""
        "她的眼里有光，那是一种经历过失去，才更懂得珍惜的温柔。"
        "婚姻从来不是童话故事的结局，而是另一段旅程的开始。"
        "在这段旅程里，我们会遇到风雨，也会看到彩虹，"
        "但很重要的前提是：我们愿意成长，愿意改变，愿意为对方成为更好的人。"
        "林姐还告诉我，二婚并不是人生的失败标签，而是一次重新认识自己的机会。"
        "第一次婚姻让她明白了什么是错的，第二次婚姻则让她学会了什么是对的。"
        "她说，如果非要给年轻人一个建议，那就是：不要把婚姻当作避风港，"
        "因为真正的安全感，从来都不是别人给的，而是自己长出来的。"
    )

    add_body(doc,
        "如果你正在婚姻中感到迷茫，不妨停下来问问自己："
        "我是不是把太多的情绪负担交给了对方？"
        "我是否真的尊重他作为独立个体的选择？"
        "我有多久没有认真经营过这段关系了？"
        "这些问题没有标准答案，但愿意去想，就已经是改变的开始。"
    )

    add_image_centered(doc, os.path.join(IMG_DIR, "img5.jpg"),
                       "夕阳下的牵手，是历经风雨后的笃定")

    add_body(doc,
        "愿我们都能在婚姻里学会爱，更学会尊重；学会依赖，更学会成长。"
        "不要等到失去了才懂得珍惜，因为有些人，一旦错过，就真的回不来了。"
        "而那些在婚姻里吃过的苦、流过的泪，终将成为我们生命里最珍贵的礼物，"
        "让我们在下一段关系里，成为更温柔、更成熟、更值得被爱的自己。"
        "人生很长，婚姻只是其中的一段路。"
        "重要的不是这条路走得有多快，而是我们是否在这段路上，成为了更好的自己。"
        "愿每一个在婚姻里迷茫的人，都能找到属于自己的答案；"
        "愿每一个经历过伤痛的人，都能在新的旅程里，遇见更温暖的风景。"
    )

    add_end_tags(doc, ["婚姻感悟", "情感成长", "二婚", "人生忠告", "亲密关系"])

    return doc


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = build_document()

    # Sanitize all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = sanitize_text(run.text)

    # Ensure length
    char_count = count_chars_in_doc(doc)
    print(f"Initial character count: {char_count}")

    # If under 2500, add more closing reflection
    if char_count < 2500:
        extra_text = (
            "婚姻的意义，从来不只是找一个人搭伙过日子。"
            "它更像是一面镜子，照出我们内心深处的渴望、恐惧和不安。"
            "每一次争吵，都是在提醒我们还有哪些伤口没有愈合；"
            "每一次和解，都是在告诉我们：爱，从来都不是一件容易的事，但它值得。"
            "林姐的故事让我明白，无论头婚还是二婚，真正决定幸福与否的，"
            "不是对方是谁，而是我们是否具备了爱人的能力。"
            "这种能力，包括管理自己的情绪，尊重对方的边界，以及日复一日的用心经营。"
            "当你拥有了这些，婚姻就不再是束缚，而是让彼此自由生长的土壤。"
        )
        extra_text = sanitize_text(extra_text)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(extra_text)
        set_run_font(run, "微软雅黑", 12, color=DARK)
        char_count = count_chars_in_doc(doc)
        print(f"After append character count: {char_count}")

    doc.save(OUT_FILE)
    print(f"DOCX saved to: {OUT_FILE}")

    # Copy images to output dir
    out_img_dir = os.path.join(OUT_DIR, "images")
    os.makedirs(out_img_dir, exist_ok=True)
    for i in range(1, 6):
        src = os.path.join(IMG_DIR, f"img{i}.jpg")
        dst = os.path.join(out_img_dir, f"img{i}.jpg")
        shutil.copy2(src, dst)
    print(f"Images copied to: {out_img_dir}")

    # ── verification ──
    doc2 = Document(OUT_FILE)
    all_text = ""
    for para in doc2.paragraphs:
        all_text += para.text + "\n"

    violations = []
    for bad in list(FORBIDDEN.keys()) + EXTRA_FORBIDDEN:
        if bad in all_text:
            violations.append(bad)

    if violations:
        print(f"WARNING: Forbidden words found: {violations}")
    else:
        print("Verification passed: no forbidden words found.")

    final_count = len(all_text.replace("\n", "").replace(" ", ""))
    print(f"Final character count (no spaces/newlines): {final_count}")
    print(f"Final character count (with spaces): {len(all_text)}")


if __name__ == "__main__":
    main()
