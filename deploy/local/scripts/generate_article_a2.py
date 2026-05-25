#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a styled docx article in 法式奶油日落风 style.
Title: 离婚律师说破真相：90%的婚姻破裂，都因为这4个字
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Style Constants ──────────────────────────────────────────────
TERRACOTTA = "D4A574"
CREAM      = "F5E6D3"
DUSTY_ROSE = "C9A9A6"
SOFT_BROWN = "8B7355"
MUTED_GOLD = "B8956A"
LIGHT_CREAM= "FDF8F0"

IMG_DIR    = "/tmp/article_a2_imgs"
OUT_DIR    = "/Users/panyong/aio_project/小程序/docs/文章/2026-05-01/1471065709"
OUT_NAME   = "离婚律师说破真相：90%的婚姻破裂，都因为这4个字.docx"
OUT_PATH   = os.path.join(OUT_DIR, OUT_NAME)

# ── Helpers ─────────────────────────────────────────────────────

def set_run_font(run, font_name="Georgia", size_pt=12, bold=False, italic=False, color_hex="000000", eastAsia="宋体"):
    """Set font with eastAsia support for CJK."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{eastAsia}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), eastAsia)

def add_paragraph(doc, text="", font_name="Georgia", size_pt=12, bold=False, italic=False,
                  color_hex="000000", alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=0, line_spacing=1.5,
                  left_indent=0, first_line_indent=0, eastAsia="宋体"):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.left_indent = Cm(left_indent)
    pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name, size_pt, bold, italic, color_hex, eastAsia)
    return p

def add_run_to_para(p, text, font_name="Georgia", size_pt=12, bold=False, italic=False,
                    color_hex="000000", eastAsia="宋体"):
    run = p.add_run(text)
    set_run_font(run, font_name, size_pt, bold, italic, color_hex, eastAsia)
    return run

def add_quote_block(doc, text, font_name="Georgia", size_pt=13, color_hex=SOFT_BROWN,
                    border_color=TERRACOTTA, bg_color=LIGHT_CREAM, eastAsia="宋体"):
    """Quote block with left border and background."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.6
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)

    # Shading background
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shd)

    # Left border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    run = p.add_run(text)
    set_run_font(run, font_name, size_pt, False, False, color_hex, eastAsia)
    return p

def add_highlight_box(doc, text, font_name="Georgia", size_pt=12, bold=True,
                      color_hex=SOFT_BROWN, bg_color=CREAM, eastAsia="宋体"):
    """Highlight box with cream background and terracotta diamond prefix."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.5
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)

    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shd)

    run1 = p.add_run("✦ ")
    set_run_font(run1, font_name, size_pt, bold, False, TERRACOTTA, eastAsia)
    run2 = p.add_run(text)
    set_run_font(run2, font_name, size_pt, bold, False, color_hex, eastAsia)
    return p

def add_chapter_marker(doc, number, title, eastAsia="宋体"):
    """Chapter marker: 01 / Title"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.3

    run1 = p.add_run(f"{number:02d}")
    set_run_font(run1, "Georgia", 24, True, False, TERRACOTTA, eastAsia)
    run2 = p.add_run(" / ")
    set_run_font(run2, "Georgia", 16, False, False, "CCCCCC", eastAsia)
    run3 = p.add_run(title)
    set_run_font(run3, "Georgia", 14, True, False, SOFT_BROWN, eastAsia)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(18)
    run = p.add_run("· · ·")
    set_run_font(run, "Georgia", 14, False, False, TERRACOTTA, "宋体")
    return p

def add_image_with_caption(doc, img_path, caption, width_inches=5.5):
    """Centered image with italic caption prefixed by ▲."""
    if not os.path.exists(img_path):
        return None
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(4)
    run = p_img.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(f"▲ {caption}")
    set_run_font(run_cap, "Georgia", 10, False, True, MUTED_GOLD, "宋体")
    return p_img

# ── Document Construction ───────────────────────────────────────

doc = Document()

# Page setup (A4)
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── English Label ──
add_paragraph(doc, "EVENING GLOW", "Georgia", 10, False, False, MUTED_GOLD,
              WD_ALIGN_PARAGRAPH.LEFT, 0, 6, 1.2)

# ── Title ──
add_paragraph(doc, "离婚律师说破真相：90%的婚姻破裂，都因为这4个字",
              "Georgia", 26, True, False, SOFT_BROWN,
              WD_ALIGN_PARAGRAPH.LEFT, 0, 8, 1.3, eastAsia="微软雅黑")

# ── Subtitle ──
add_paragraph(doc, "一位从业十五年的家事律师，在法庭上见证了太多爱情的终结",
              "Georgia", 18, False, False, TERRACOTTA,
              WD_ALIGN_PARAGRAPH.LEFT, 0, 8, 1.3, eastAsia="微软雅黑")

# ── Tagline ──
add_paragraph(doc, "婚姻不是爱情的坟墓，\"理所当然\"才是",
              "Georgia", 13, False, True, DUSTY_ROSE,
              WD_ALIGN_PARAGRAPH.LEFT, 0, 18, 1.3, eastAsia="微软雅黑")

# ── Image 1 ──
add_image_with_caption(doc, os.path.join(IMG_DIR, "img1.jpg"),
                       "黄昏时分，一对夫妻背对背坐在长椅上")

# ── Opening ──
add_paragraph(doc,
    "从业十五年，我经手过上千起离婚案件。每一次坐在法庭的旁听席上，看着曾经相爱的两个人针锋相对，"
    "我都会想起一个问题：到底是什么，把一段婚姻推向了无法挽回的深渊？",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "出轨？家暴？经济纠纷？这些确实都是离婚的理由。但如果你问我，"
    "绝大多数婚姻走向破裂的真正原因，我会告诉你四个字——",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_highlight_box(doc, "理所当然。")

add_paragraph(doc,
    "这四个字，看似平淡无奇，却像一把钝刀，日复一日地割开感情的纽带。"
    "它不是一次性的伤害，而是日积月累的冷漠；它不是轰轰烈烈的争吵，而是悄无声息的疏离。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_divider(doc)

# ── Chapter 1 ──
add_chapter_marker(doc, 1, "理所当然，是婚姻里的慢性毒药")

add_paragraph(doc,
    "我有一个当事人，结婚十二年，丈夫是一家公司的中层管理，她是全职妈妈。"
    "来咨询离婚那天，她穿着朴素的针织衫，眼圈泛红，却努力保持着体面。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "她说：\"我以为，我为这个家付出的一切，他都看在眼里。\""
    "\"我每天五点起床做早餐，送孩子上学，打扫房子，准备晚饭，辅导作业……"
    "我以为这些他都懂，我以为他知道我有多累。\"",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "\"结果呢？\"她苦笑了一下，\"他觉得这一切都是理所当然的。"
    "理所当然地有干净的衣服穿，理所当然地有热腾腾的饭菜，"
    "理所当然地孩子成绩优异、家里井井有条。他甚至连一句谢谢都懒得说。\"",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_quote_block(doc,
    "\"理所当然\"这四个字，是婚姻里最隐蔽的杀手。它让付出变得廉价，让感激变得多余，"
    "让原本应该被珍视的一切，变成了对方眼中\"本该如此\"的义务。")

add_paragraph(doc,
    "心理学上有一个概念叫\"情感账户\"。每一段亲密关系，都像是一个共同的银行账户。"
    "你每一次表达感谢、每一个温暖的拥抱、每一句\"辛苦了\"，都是在往这个账户里存款。"
    "而每一次\"理所当然\"的索取、每一次视而不见的冷漠，都是在透支。"
    "当账户余额归零，甚至变成负数的时候，这段关系也就走到了尽头。"
    "很多人直到感情破裂，才恍然大悟：原来那些被我忽略的日常，才是婚姻里最珍贵的部分。"
    "一个清晨的早安吻，一顿用心准备的晚餐，一次耐心的倾听——"
    "这些看似微不足道的瞬间，恰恰是维系感情的关键纽带。"
    "可惜的是，\"理所当然\"的心态让我们对这些美好视而不见，"
    "直到对方心灰意冷，选择离开，我们才追悔莫及。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

# ── Image 2 ──
add_image_with_caption(doc, os.path.join(IMG_DIR, "img2.jpg"),
                       "空荡荡的餐桌，只剩下一副碗筷")

add_divider(doc)

# ── Chapter 2 ──
add_chapter_marker(doc, 2, "从\"谢谢你\"到\"随便你\"，只需要三年")

add_paragraph(doc,
    "我统计过自己经手的案件，发现一个很惊人的规律："
    "绝大多数婚姻危机，并不是在某一个瞬间爆发的，而是在日常琐碎中，一点点被蚕食。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "有一位男当事人让我印象很深。他是做IT的，收入不错，性格也温和。"
    "妻子提出离婚时，他完全无法理解：\"我又没出轨，又没家暴，"
    "每个月工资都上交，她到底还有什么不满意的？\"",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "后来我在调解过程中了解到，结婚第三年，妻子生日那天，她做了一桌子菜等他回来。"
    "他加班到九点，进门第一句话是：\"怎么还没收拾？\""
    "然后径直去洗澡了。妻子一个人坐在餐桌前，看着蜡烛燃尽，蛋糕上的奶油融化成一滩。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "\"我不是非要他买什么礼物，\"妻子说，\"我只是希望，他能记得。"
    "哪怕只是一句生日快乐，哪怕只是一个拥抱。\"",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_highlight_box(doc, "婚姻里的失望，从来不是来自惊天动地的大事，而是来自那些被忽略的小事。")

add_paragraph(doc,
    "从\"谢谢你\"到\"随便你\"，从\"我来吧\"到\"你自己弄\"，"
    "从\"今天过得怎么样\"到沉默的晚餐——这段距离，往往只需要三年。"
    "而走到这一步的夫妻，通常还没有意识到问题的严重性。"
    "他们以为只是\"感情淡了\"，却不知道，\"理所当然\"已经在他们之间筑起了一堵高墙。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

# ── Image 3 ──
add_image_with_caption(doc, os.path.join(IMG_DIR, "img3.jpg"),
                       "夕阳下独自散步的身影")

add_divider(doc)

# ── Chapter 3 ──
add_chapter_marker(doc, 3, "为什么我们总是对陌生人客气，对爱人苛刻？")

add_paragraph(doc,
    "这是我在咨询室里，经常问当事人的一个问题。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "你会对同事说\"麻烦你了\"，会对服务员说\"谢谢\"，"
    "会对快递小哥说\"辛苦了\"——可是对那个每天为你做饭、洗衣、带孩子的人，"
    "你却连一句简单的感谢都觉得多余。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "为什么会这样？因为\"理所当然\"的心理机制在作祟。"
    "当一个人的付出变成了日常，我们的大脑就会自动将其归类为\"背景噪音\"，"
    "从而失去感知和回应的能力。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_quote_block(doc,
    "我们对陌生人的善意心存感激，是因为那是\"额外\"的；"
    "我们对伴侣的付出视而不见，是因为我们误以为那是\"应该\"的。")

add_paragraph(doc,
    "更可怕的是，\"理所当然\"往往是双向的。"
    "丈夫觉得妻子照顾家庭是理所当然，妻子也觉得丈夫赚钱养家是理所当然。"
    "两个人都在付出，两个人都觉得自己的付出被忽视，"
    "于是怨气越积越深，直到某一天，一件微不足道的小事，成了压垮骆驼的最后一根稻草。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "我见过太多这样的案例：离婚的原因，表面上是\"性格不合\"\"三观不一致\"，"
    "但深挖下去，根源都是同一个——两个人都觉得自己在单向付出，"
    "而对方\"理所当然\"地享受着一切，从不表达感激。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

# ── Image 4 ──
add_image_with_caption(doc, os.path.join(IMG_DIR, "img4.jpg"),
                       "一对夫妻隔着玻璃门，各自望向不同的方向")

add_divider(doc)

# ── Chapter 4 ──
add_chapter_marker(doc, 4, "打破\"理所当然\"，从看见开始")

add_paragraph(doc,
    "说了这么多，并不是想让大家对婚姻感到绝望。"
    "恰恰相反，我想告诉你的是：\"理所当然\"虽然可怕，但它并非不可战胜。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "在我经手的案件中，也有相当一部分夫妻，在濒临破裂的边缘选择了回头。"
    "他们中的很多人告诉我，真正挽救婚姻的，往往不是昂贵的礼物或浪漫的旅行，"
    "而是一些极其微小的改变。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_highlight_box(doc, "看见对方的付出，并表达出来。这是打破\"理所当然\"的第一步，也是最关键的一步。")

add_paragraph(doc,
    "比如，丈夫下班回来，妻子递上一杯温水，他说一句\"谢谢，辛苦了\"；"
    "妻子做完晚饭，丈夫主动收拾碗筷，她说一句\"有你在真好\"。"
    "这些话语本身不值钱，但它们传递的信息是：我看见了你，我感激你，我珍惜你。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "有一位复婚的当事人跟我说过一句话，我至今记得很清楚："
    "\"我们离婚的那一年，我才意识到，她每天早起给我做的那碗粥，"
    "不是理所当然的。那是她爱我的一种方式。\"",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "可惜的是，很多人要等到失去之后，才明白这个道理。"
    "而婚姻最残酷的地方就在于：有些伤害一旦造成，就再也无法弥补。"
    "所以，如果你还在这段关系里，如果你还爱着对方，"
    "请从今天开始，把\"理所当然\"从你们的字典里删掉。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

# ── Image 5 ──
add_image_with_caption(doc, os.path.join(IMG_DIR, "img5.jpg"),
                       "黄昏中牵手的老夫妻，背影温暖而坚定")

add_divider(doc)

# ── Closing ──
add_paragraph(doc,
    "作为一名离婚律师，我比大多数人更清楚婚姻的脆弱。"
    "但我也比大多数人更相信，真正让婚姻长久的，不是激情，不是誓言，"
    "而是日复一日的珍惜与感恩。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_paragraph(doc,
    "别让\"理所当然\"，偷走你的幸福。"
    "因为在这个世界上，没有什么是真正理所当然的——"
    "包括那个愿意陪你走过漫长岁月的人。",
    "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.LEFT,
    0, 8, 1.8, 0, 0.74, "宋体")

add_quote_block(doc,
    "婚姻不是找到一个完美的人，而是学会用完美的眼光，欣赏一个不完美的人。"
    "而这份欣赏的起点，就是不再把对方的一切，都当作理所当然。")

# ── End Tags ──
add_paragraph(doc, "", "Georgia", 12, False, False, "333333", WD_ALIGN_PARAGRAPH.CENTER, 12, 4, 1.2)
add_paragraph(doc, "#婚姻感悟", "Georgia", 11, False, False, MUTED_GOLD, WD_ALIGN_PARAGRAPH.CENTER, 0, 4, 1.2)
add_paragraph(doc, "#情感成长", "Georgia", 11, False, False, MUTED_GOLD, WD_ALIGN_PARAGRAPH.CENTER, 0, 4, 1.2)
add_paragraph(doc, "#亲密关系", "Georgia", 11, False, False, MUTED_GOLD, WD_ALIGN_PARAGRAPH.CENTER, 0, 4, 1.2)
add_paragraph(doc, "# EveningGlow", "Georgia", 10, False, True, DUSTY_ROSE, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.2)

# ── Save ──
doc.save(OUT_PATH)
print(f"Document saved to: {OUT_PATH}")

# ── Copy images to output dir ──
for f in os.listdir(IMG_DIR):
    src = os.path.join(IMG_DIR, f)
    dst = os.path.join(OUT_DIR, f)
    if os.path.isfile(src):
        shutil.copy2(src, dst)
        print(f"Copied image: {f}")

# ── Character count ──
full_text = ""
for para in doc.paragraphs:
    full_text += para.text
char_count = len(full_text.replace(" ", "").replace("\n", "").replace("\t", ""))
print(f"Total character count (no spaces): {char_count}")
