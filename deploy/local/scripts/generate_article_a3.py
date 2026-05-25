#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a styled docx article in 清新马卡龙风 style.
"""

import os
import re
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# Color definitions
# ============================================================
COLOR_PINK = "F4A6B3"
COLOR_TEAL = "5B8C85"
COLOR_LIGHT_TEAL = "7AAFA5"
COLOR_GRAY = "A0AEC0"
COLOR_DARK_GRAY = "718096"
COLOR_PURPLE_BORDER = "C9B1FF"
COLOR_LIGHT_PURPLE_BG = "FAF8FF"
COLOR_MINT_BG = "F0FAF7"
COLOR_CREAM_BG = "FFF8F0"
COLOR_LIGHT_BORDER = "E2E8F0"
COLOR_DARK_TEXT = "4A5568"

# ============================================================
# Forbidden words and replacements
# ============================================================
FORBIDDEN_REPLACEMENTS = {
    "永久": "长期",
    "关注": "留意",
    "最新": "新近",
    "最大": "很大",
    "群": "圈子",
    "第一": "领先",
    "唯一": "少有的",
    "顶级": "优质",
    "最高级": "更高级",
    "万能": "多用途",
    "100%": "绝大多数",
    "绝对": "相当",
    "最好": "很好",
    "最佳": "很出色",
    "最强": "很强",
    "首选": "优选",
    "全网最低": "价格优惠",
    "最便宜": "性价比高",
    "国家级": "省级以上",
    "首家": "率先",
    "独家": "特有",
    "史上最": "非常",
    "特效": "效果明显",
    "震惊": "令人关注",
    "跳楼价": "超值价",
    "亏本卖": "薄利多销",
    "错过再等一年": "限时优惠",
    "保证治愈": "有望改善",
    "药到病除": "有助缓解",
    "治愈": "改善",
    "根治": "从根本上改善",
    "无副作用": "成分温和",
    "疗效": "效果",
    "秘方": "传统方法",
    "偏方": "传统方法",
    "神医": "资深医师",
    "神药": "有效药物",
    "无效退款": "可咨询售后",
    "抗癌": "辅助调理",
    "增高": "促进发育",
    "一吃就瘦": "配合运动",
    "减肥神药": "辅助产品",
    "延寿": "健康养生",
    "几天见效": "坚持服用",
    "暴富": "财富增长",
    "保本保息": "稳健型",
    "稳赚不赔": "收益相对稳定",
    "保本": "稳健型",
    "稳赚": "收益相对稳定",
    "零风险": "风险可控",
    "原始股": "股权投资",
    "内幕消息": "市场分析",
    "配资": "融资服务",
    "收益率100%": "预期收益",
    "高回报": "预期收益较好",
    "假一赔十": "正品保障",
    "免费送": "赠品活动",
    "烟草": "其他商品",
    "爆炸性": "重要",
    "突发": "刚刚",
    "重磅": "重要",
    "速看": "值得关注",
    "躺赚": "被动收入",
    "限时免费": "免费体验",
    "零元购": "免费试用",
    "全网首发": "率先发布",
    "白菜价": "性价比高",
    "先到先得": "数量有限",
    "转发此文章": "欢迎分享",
    "点击关注": "欢迎关注",
    "最后机会": "机会难得",
    "算命": "运势分析",
    "改运": "积极心态",
    "约炮": "交友",
    "一夜情": "短期关系",
    "赌博": "娱乐",
    "彩票预测": "彩票分析",
    "时时彩": "彩票",
    "外挂": "辅助工具",
    "走私": "跨境贸易",
    "破解版": "正版授权",
    "盗版": "非正版",
    "VPN": "网络工具",
    "翻墙": "跨境访问",
    "枪支": "器械",
    "代购": "海外购",
    "占卜": "性格测试",
    "八字": "出生日期",
    "风水": "环境布局",
    "开光": "仪式",
    "写真": "照片",
    "命格": "性格特点",
    "转运": "好运气",
    "荐股": "投资建议",
    "电子烟": "雾化器",
    "百分百": "绝大多数",
}

FORBIDDEN_PHRASES = [
    "分享到朋友圈', '不转不是中国人', '不转发死全家', '转发好运', '福利姬', '无码', '有码"
]


def set_run_font(run, font_name="微软雅黑", size_pt=12, bold=False, italic=False, color_hex=None):
    """Set font properties for a run with eastAsia support."""
    font = run.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color_hex:
        font.color.rgb = RGBColor.from_string(color_hex)
    # Set eastAsia font for CJK characters
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_paragraph_with_style(doc, text, font_name="微软雅黑", size_pt=12, bold=False, italic=False,
                              color_hex=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              space_before=0, space_after=0, line_spacing=1.5,
                              left_indent=0, right_indent=0,
                              border_bottom=False, border_bottom_color=None,
                              border_top=False, border_top_color=None,
                              shading_color=None):
    """Add a paragraph with comprehensive styling."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if left_indent:
        pf.left_indent = Pt(left_indent)
    if right_indent:
        pf.right_indent = Pt(right_indent)

    run = p.add_run(text)
    set_run_font(run, font_name, size_pt, bold, italic, color_hex)

    # Paragraph borders via XML
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)

    if border_bottom or border_top or shading_color:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
            pPr.append(pBdr)

        if border_bottom:
            bottom = parse_xml(
                f'<w:bottom {nsdecls("w")} w:val="single" w:sz="4" w:space="1" w:color="{border_bottom_color or COLOR_LIGHT_BORDER}"/>'
            )
            # Remove existing bottom if any
            existing = pBdr.find(qn('w:bottom'))
            if existing is not None:
                pBdr.remove(existing)
            pBdr.append(bottom)

        if border_top:
            top = parse_xml(
                f'<w:top {nsdecls("w")} w:val="single" w:sz="4" w:space="1" w:color="{border_top_color or COLOR_LIGHT_BORDER}"/>'
            )
            existing = pBdr.find(qn('w:top'))
            if existing is not None:
                pBdr.remove(existing)
            pBdr.append(top)

    if shading_color:
        shd = pPr.find(qn('w:shd'))
        if shd is None:
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" w:val="clear"/>')
            pPr.append(shd)
        else:
            shd.set(qn('w:fill'), shading_color)
            shd.set(qn('w:val'), 'clear')

    return p


def add_empty_paragraph(doc, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_image_with_caption(doc, img_path, width_inches=5.5, caption_text=""):
    """Add a centered image with optional caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    if caption_text:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(f"▲ {caption_text}")
        set_run_font(cap_run, "Georgia", 10, False, True, COLOR_GRAY)
        cap_p.paragraph_format.space_after = Pt(12)


def add_divider(doc):
    """Add a centered divider: · · · in pink."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("·  ·  ·")
    set_run_font(run, "Georgia", 14, False, False, COLOR_PINK)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)


def add_quote_block(doc, text):
    """Add a quote block with top/bottom borders and mint background."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.left_indent = Pt(24)
    pf.right_indent = Pt(24)
    pf.line_spacing = 1.6

    run = p.add_run(text)
    set_run_font(run, "微软雅黑", 13, False, False, COLOR_DARK_GRAY)

    # Borders and shading via XML
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)

    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
        pPr.append(pBdr)

    for side, sz in [("top", "6"), ("bottom", "6")]:
        elem = parse_xml(
            f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="1" w:color="{COLOR_LIGHT_BORDER}"/>'
        )
        existing = pBdr.find(qn(f'w:{side}'))
        if existing is not None:
            pBdr.remove(existing)
        pBdr.append(elem)

    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_MINT_BG}" w:val="clear"/>')
        pPr.append(shd)
    else:
        shd.set(qn('w:fill'), COLOR_MINT_BG)
        shd.set(qn('w:val'), 'clear')


def add_highlight_box1(doc, text):
    """Add highlight box 1: purple border, light purple bg, ✿ pink prefix."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.left_indent = Pt(12)
    pf.right_indent = Pt(12)
    pf.line_spacing = 1.6

    run = p.add_run(f"✿ {text}")
    set_run_font(run, "微软雅黑", 12, True, False, COLOR_DARK_TEXT)

    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)

    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
        pPr.append(pBdr)

    for side in ["top", "bottom", "left", "right"]:
        elem = parse_xml(
            f'<w:{side} {nsdecls("w")} w:val="single" w:sz="6" w:space="1" w:color="{COLOR_PURPLE_BORDER}"/>'
        )
        existing = pBdr.find(qn(f'w:{side}'))
        if existing is not None:
            pBdr.remove(existing)
        pBdr.append(elem)

    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_LIGHT_PURPLE_BG}" w:val="clear"/>')
        pPr.append(shd)
    else:
        shd.set(qn('w:fill'), COLOR_LIGHT_PURPLE_BG)
        shd.set(qn('w:val'), 'clear')


def add_highlight_box2(doc, text):
    """Add highlight box 2: pink border, cream bg, bold teal text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.left_indent = Pt(12)
    pf.right_indent = Pt(12)
    pf.line_spacing = 1.6

    run = p.add_run(text)
    set_run_font(run, "微软雅黑", 12, True, False, COLOR_TEAL)

    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)

    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
        pPr.append(pBdr)

    for side in ["top", "bottom", "left", "right"]:
        elem = parse_xml(
            f'<w:{side} {nsdecls("w")} w:val="single" w:sz="6" w:space="1" w:color="{COLOR_PINK}"/>'
        )
        existing = pBdr.find(qn(f'w:{side}'))
        if existing is not None:
            pBdr.remove(existing)
        pBdr.append(elem)

    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_CREAM_BG}" w:val="clear"/>')
        pPr.append(shd)
    else:
        shd.set(qn('w:fill'), COLOR_CREAM_BG)
        shd.set(qn('w:val'), 'clear')


def add_chapter_marker(doc, number, title):
    """Add chapter marker: number in pink, / in gray, title in teal."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(20)
    pf.space_after = Pt(8)

    run1 = p.add_run(str(number))
    set_run_font(run1, "Georgia", 24, True, False, COLOR_PINK)

    run2 = p.add_run(" / ")
    set_run_font(run2, "Georgia", 16, False, False, COLOR_GRAY)

    run3 = p.add_run(title)
    set_run_font(run3, "微软雅黑", 14, True, False, COLOR_TEAL)

    # Section bottom border
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
        pPr.append(pBdr)
    bottom = parse_xml(
        f'<w:bottom {nsdecls("w")} w:val="single" w:sz="4" w:space="1" w:color="{COLOR_LIGHT_BORDER}"/>'
    )
    existing = pBdr.find(qn('w:bottom'))
    if existing is not None:
        pBdr.remove(existing)
    pBdr.append(bottom)


def add_body_text(doc, text, highlight_phrases=None):
    """Add body text with optional bold pink highlights."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.6

    if highlight_phrases:
        remaining = text
        # Sort phrases by length descending to avoid partial matches
        sorted_phrases = sorted(highlight_phrases, key=len, reverse=True)
        while remaining:
            found = False
            for phrase in sorted_phrases:
                idx = remaining.find(phrase)
                if idx != -1:
                    if idx > 0:
                        run = p.add_run(remaining[:idx])
                        set_run_font(run, "微软雅黑", 12, False, False, None)
                    run = p.add_run(phrase)
                    set_run_font(run, "微软雅黑", 12, True, False, COLOR_PINK)
                    remaining = remaining[idx + len(phrase):]
                    found = True
                    break
            if not found:
                run = p.add_run(remaining)
                set_run_font(run, "微软雅黑", 12, False, False, None)
                break
    else:
        run = p.add_run(text)
        set_run_font(run, "微软雅黑", 12, False, False, None)


def add_end_tags(doc, tags):
    """Add centered end tags with #xxx format."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(20)
    pf.space_after = Pt(8)
    run = p.add_run(" ".join(f"#{t}" for t in tags))
    set_run_font(run, "微软雅黑", 11, False, False, COLOR_GRAY)


def set_page_default_font(doc, font_name="微软雅黑", size_pt=12):
    """Set default document font via styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def check_forbidden_words(text):
    """Check text for forbidden words and phrases."""
    found = []
    for phrase in FORBIDDEN_PHRASES:
        if phrase in text:
            found.append(phrase)
    for word in FORBIDDEN_REPLACEMENTS:
        if word in text:
            found.append(word)
    return found


def sanitize_text(text):
    """Replace forbidden words in text."""
    for phrase in FORBIDDEN_PHRASES:
        text = text.replace(phrase, "")
    for word, replacement in FORBIDDEN_REPLACEMENTS.items():
        text = text.replace(word, replacement)
    # Also remove standalone 最 where possible (simple heuristic)
    text = re.sub(r'(?<![\u4e00-\u9fff])最(?![\u4e00-\u9fff])', '很', text)
    return text


def build_article():
    doc = Document()
    set_page_default_font(doc, "微软雅黑", 12)

    # Set narrow margins for better layout
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== TOP BORDER (pink bottom border on first empty paragraph) ==========
    top_p = doc.add_paragraph()
    top_p.paragraph_format.space_after = Pt(2)
    pPr = top_p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        top_p._element.insert(0, pPr)
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
        pPr.append(pBdr)
    bottom = parse_xml(
        f'<w:bottom {nsdecls("w")} w:val="single" w:sz="12" w:space="1" w:color="{COLOR_PINK}"/>'
    )
    pBdr.append(bottom)

    add_empty_paragraph(doc, 4)

    # ========== ENGLISH LABEL ==========
    add_paragraph_with_style(doc, "SWEET MOMENTS", "Georgia", 10, False, False,
                              COLOR_PINK, WD_ALIGN_PARAGRAPH.LEFT, 0, 4)

    # ========== TITLE ==========
    title_text = "为什么越懂事的女人越不被珍惜？这是我花了5年婚姻才懂的道理"
    add_paragraph_with_style(doc, title_text, "微软雅黑", 28, True, False,
                              COLOR_TEAL, WD_ALIGN_PARAGRAPH.LEFT, 0, 8)

    # ========== SUBTITLE ==========
    add_paragraph_with_style(doc, "在爱里，懂事不是勋章，而是枷锁", "微软雅黑", 18, True, False,
                              COLOR_LIGHT_TEAL, WD_ALIGN_PARAGRAPH.LEFT, 0, 6)

    # ========== TAGLINE ==========
    add_paragraph_with_style(doc, "写给每一个在婚姻里学会了沉默，却弄丢了自己的女人", "微软雅黑", 13, False, True,
                              COLOR_GRAY, WD_ALIGN_PARAGRAPH.LEFT, 0, 12)

    add_empty_paragraph(doc, 4)

    # ========== IMAGE 1 ==========
    img1 = "/tmp/article_a3_imgs/img1.jpg"
    add_image_with_caption(doc, img1, 5.5, "婚姻不是一个人的独角戏")

    add_empty_paragraph(doc, 6)

    # ========== INTRO ==========
    intro_text = (
        "结婚第五年，我终于在深夜的厨房里哭了出来。"
        "不是因为吵架，不是因为委屈，而是因为我忽然意识到——"
        "这五年来，我把自己活成了一个'透明人'。"
        "我体谅他工作辛苦，从不让他做家务；"
        "我理解他经济压力大，从不要求礼物和惊喜；"
        "我心疼他应酬多，从不查岗不追问。"
        "我以为这样的懂事，会换来珍惜。"
        "可现实是，我的体贴成了理所当然，我的沉默成了默认许可。"
        "他不再问我今天过得好不好，不再留意我换了新发型，"
        "甚至在我生病发烧的时候，也只是淡淡地说一句：多喝热水。"
    )
    add_body_text(doc, intro_text, ["透明人", "理所当然", "默认许可"])

    add_divider(doc)

    # ========== CHAPTER 1 ==========
    add_chapter_marker(doc, "01", "懂事，是婚姻里最危险的品质")

    c1_p1 = (
        "小时候，父母总说：你要懂事，要听话，要体谅别人。"
        "于是我把'懂事'刻进了骨子里。"
        "谈恋爱的时候，我不吵不闹，他忘了纪念日，我说没关系；"
        "他约会迟到两小时，我说你忙我理解。"
        "结婚后，我更是把'懂事'发挥到了极致。"
        "家里的大小事务，我一个人扛；"
        "婆媳之间的小摩擦，我一个人忍；"
        "孩子半夜哭闹，我一个人哄。"
        "我以为我是在经营一个温暖的家，"
        "却不知道，我的'懂事'正在一点点抹去自己的存在。"
    )
    add_body_text(doc, c1_p1, ["懂事", "抹去自己的存在"])

    add_quote_block(doc, (
        "心理学家说：在亲密关系里，过度付出的一方，"
        "往往会在不知不觉中降低自己的情感价值。"
        "当你的付出没有边界，对方就会失去感恩的能力。"
    ))

    c1_p2 = (
        "我开始回想，上一次他认真看着我的眼睛说'谢谢你'，是什么时候？"
        "上一次他主动问我'你累不累'，又是什么时候？"
        "答案让我心酸——我已经想不起来了。"
        "不是他变了，是我的'懂事'让他习惯了忽略。"
        "他习惯了饭菜自动上桌，习惯了衣服自动洗净叠好，"
        "习惯了家里永远整洁、孩子永远安静。"
        "他以为这一切都是自然而然发生的，"
        "却忘了，这背后是一个女人在用尽全力。"
    )
    add_body_text(doc, c1_p2, ["习惯了忽略", "用尽全力"])

    add_empty_paragraph(doc, 4)
    img2 = "/tmp/article_a3_imgs/img2.jpg"
    add_image_with_caption(doc, img2, 5.5, "一个人的付出，撑不起两个人的婚姻")

    add_empty_paragraph(doc, 6)

    # ========== CHAPTER 2 ==========
    add_chapter_marker(doc, "02", "你的体谅，成了他的免责金牌")

    c2_p1 = (
        "有一次，他答应周末带孩子去公园，结果临时被朋友叫去打球。"
        "我笑着说：去吧，孩子我带去就行。"
        "他如释重负地走了，留下我和孩子在烈日下排了一个小时的队。"
        "那天晚上，他回来兴致勃勃地讲球赛，"
        "却没有问一句：你们今天玩得开心吗？"
        "这样的场景，在我们的婚姻里反复上演。"
        "我体谅他的工作压力，所以从不抱怨他加班；"
        "我体谅他的社交需求，所以从不干涉他聚会；"
        "我体谅他的疲惫，所以从不要求他分担家务。"
        "可我的体谅，换来的不是感激，而是越来越多的'理所当然'。"
    )
    add_body_text(doc, c2_p1, ["如释重负", "理所当然"])

    add_highlight_box1(doc, "真正的爱不是不计较，而是计较了还愿意一起解决。"
                           "你的每一次'没关系'，都在教会他：你的感受不重要。")

    c2_p2 = (
        "朋友小雯曾经跟我说过一句话，当时我还不以为然："
        "女人不能太懂事，懂事的女人没人疼。"
        "现在我懂了。"
        "因为当你总是说'没关系'的时候，"
        "对方就真的以为，一切都没关系。"
        "你的委屈、你的疲惫、你的渴望，"
        "在一次次的'懂事'中被深深掩埋。"
        "直到有一天，你发现自己已经不会表达需求了，"
        "因为你知道，即使说了，也不会有人听。"
    )
    add_body_text(doc, c2_p2, ["没关系", "深深掩埋"])

    add_empty_paragraph(doc, 4)
    img3 = "/tmp/article_a3_imgs/img3.jpg"
    add_image_with_caption(doc, img3, 5.5, "沉默不是金，是婚姻里的慢性毒药")

    add_empty_paragraph(doc, 6)

    # ========== CHAPTER 3 ==========
    add_chapter_marker(doc, "03", "我学会了'不懂事'，婚姻反而变好了")

    c3_p1 = (
        "转折发生在我三十岁那年的生日。"
        "我提前一周告诉他，我想去那家新开的餐厅吃饭。"
        "他答应了。"
        "可生日那天，他临时说公司有事，要晚点回来。"
        "换作以前，我会说：没事，工作要紧。"
        "但那天，我突然不想懂事了。"
        "我说：今天是我的生日，这家餐厅我预约了很久，"
        "你可以处理完事情后来找我，但我不会取消预约。"
        "他愣了一下，然后说了声'好'。"
        "那天晚上，他迟到了一个小时，但还是来了。"
        "虽然不完美，但那是我结婚以来，头一次感受到："
        "我的需求，是被看见的。"
    )
    add_body_text(doc, c3_p1, ["不想懂事了", "被看见的"])

    add_highlight_box2(doc, "懂事是一种选择，但不应该成为仅有的选项。"
                           "你有权利表达不满，有权利要求陪伴，有权利被优先考虑。")

    c3_p2 = (
        "从那以后，我开始学着'不懂事'。"
        "累了就说累了，不想做饭就点外卖，"
        "想要礼物就直说，不高兴了就表达出来。"
        "一开始他很诧异，甚至有些不适应。"
        "但慢慢地，他开始留意我的情绪变化，"
        "开始主动问我'今天过得怎么样'，"
        "开始在我加班的时候发来关心的消息。"
        "原来，不是他不会爱，是我从来没有给过他爱我的机会。"
        "当我总是把'我没事'挂在嘴边的时候，"
        "他怎么可能知道，我其实有很多事？"
    )
    add_body_text(doc, c3_p2, ["不懂事", "不会爱", "爱我的机会"])

    add_empty_paragraph(doc, 4)
    img4 = "/tmp/article_a3_imgs/img4.jpg"
    add_image_with_caption(doc, img4, 5.5, "表达需求，是爱的另一种语言")

    add_empty_paragraph(doc, 6)

    # ========== CHAPTER 4 ==========
    add_chapter_marker(doc, "04", "写给每一个正在经历的女人")

    c4_p1 = (
        "如果你正在读这篇文章，并且觉得每一句话都像是在说自己，"
        "那么我想告诉你：停下来，好好看看自己。"
        "你是不是也已经很久没有为自己买过一件像样的衣服了？"
        "你是不是也总是把好的东西留给他和孩子，自己随便对付？"
        "你是不是也已经忘记了，上一次被宠爱、被呵护是什么感觉？"
        "亲爱的，你的善良和体贴很珍贵，"
        "但请不要把它们浪费在不懂珍惜的人身上。"
        "即使那个人是你的丈夫。"
    )
    add_body_text(doc, c4_p1, ["停下来", "很珍贵", "不懂珍惜"])

    add_quote_block(doc, (
        "婚姻不是牺牲自我来成全对方，而是两个完整的人，"
        "选择一起变得更好。"
        "如果你在这段关系里越来越渺小，那不是爱情，那是消耗。"
    ))

    c4_p2 = (
        "我开始重新拾起自己的爱好。"
        "周末去上瑜伽课，偶尔和闺蜜喝下午茶，"
        "买了自己喜欢的书，在睡前读上几页。"
        "我不再把全部的精力都放在家庭上，"
        "而是留出一部分，好好照顾自己。"
        "奇妙的是，当我开始爱自己，"
        "他似乎也更爱我了。"
        "他会在我练完瑜伽后递来一杯水，"
        "会在我看书的时候安静地陪在身边，"
        "会主动提出周末他来带孩子，让我休息。"
        "原来，爱从来不是求来的，而是吸引来的。"
        "当你把自己活成一道光，"
        "自然会有人愿意靠近你、温暖你。"
    )
    add_body_text(doc, c4_p2, ["爱自己", "吸引来的", "一道光"])

    add_empty_paragraph(doc, 4)
    img5 = "/tmp/article_a3_imgs/img5.jpg"
    add_image_with_caption(doc, img5, 5.5, "当你爱自己，世界才会来爱你")

    add_empty_paragraph(doc, 6)

    # ========== CONCLUSION ==========
    add_chapter_marker(doc, "05", "懂事可以，但别弄丢了自己")

    conclusion = (
        "五年的婚姻，让我明白了一个很深刻的道理："
        "懂事是一种美德，但过度的懂事是一种自我伤害。"
        "你可以体贴，但不要无底线索取；"
        "你可以包容，但不要没有原则退让；"
        "你可以付出，但不要忘记自己也需要被爱。"
        "真正健康的婚姻，不是一个人拼命懂事，另一个人安心享受。"
        "而是两个人都能坦诚地表达需求，"
        "都能在对方面前做真实的自己。"
        "你可以撒娇，可以任性，可以说'我今天不开心'。"
        "因为真正爱你的人，不会嫌你麻烦，"
        "他会心疼你的委屈，珍惜你的坦诚。"
        "所以，从今天起，做一个'适度懂事'的女人吧。"
        "懂事给值得的人，任性留给爱自己。"
        "你的温柔要有锋芒，你的善良要有底线。"
        "只有这样，你才能在婚姻里，既温暖了别人，也照亮了自己。"
        "记住，你值得被好好对待，不是因为你付出了多少，而是因为你是独一无二的你。"
        "不要让懂事成为束缚你的枷锁，而要让它成为你选择爱的方式。"
        "当你学会在爱里保持自我，你的婚姻才会真正焕发光彩。"
        "这不是自私，这是对自己负责，也是对这段关系负责。"
        "愿你在婚姻里，既能温柔待人，也能被温柔以待。"
    )
    add_body_text(doc, conclusion, ["自我伤害", "真实的自己", "适度懂事", "温暖了别人，也照亮了自己"])

    add_divider(doc)

    add_paragraph_with_style(doc, (
        "愿每一个在婚姻里努力过的你，都能被温柔以待。"
        "不是因为你懂事，而是因为你是你。"
    ), "微软雅黑", 13, False, False, COLOR_DARK_GRAY, WD_ALIGN_PARAGRAPH.CENTER, 8, 8)

    # ========== END TAGS ==========
    add_end_tags(doc, ["婚姻感悟", "女性成长", "情感共鸣", "自我成长", "爱情观"])

    return doc


def main():
    out_dir = "/Users/panyong/aio_project/小程序/docs/文章/2026-05-01/朱朱"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "为什么越懂事的女人越不被珍惜？这是我花了5年婚姻才懂的道理.docx")

    doc = build_article()
    doc.save(out_path)
    print(f"Docx saved to: {out_path}")

    # Copy images to output directory
    img_out_dir = os.path.join(out_dir, "images")
    os.makedirs(img_out_dir, exist_ok=True)
    for i in range(1, 6):
        src = f"/tmp/article_a3_imgs/img{i}.jpg"
        dst = os.path.join(img_out_dir, f"img{i}.jpg")
        shutil.copy2(src, dst)
    print(f"Images copied to: {img_out_dir}")

    # Read back and verify no forbidden words
    doc2 = Document(out_path)
    full_text = "\n".join([p.text for p in doc2.paragraphs])
    # Also check tables if any
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    forbidden_found = check_forbidden_words(full_text)
    if forbidden_found:
        print(f"WARNING: Forbidden words found: {forbidden_found}")
        # Fix by rebuilding with sanitized text - for now just report
    else:
        print("No forbidden words found. Document is clean.")

    char_count = len(full_text.replace(" ", "").replace("\n", ""))
    print(f"Total character count (no spaces): {char_count}")


if __name__ == "__main__":
    main()
