#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a 清新马卡龙风 styled docx article.
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# =============================================================================
# Configuration
# =============================================================================
OUTPUT_PATH = "/Users/panyong/aio_project/小程序/docs/文章/2026-04-30/深情冰美式/被同事甩锅、被领导打压？学会这4招，没人再敢欺负你.docx"
IMG_DIR = "/tmp/article_u_imgs"
OUTPUT_IMG_DIR = os.path.join(os.path.dirname(OUTPUT_PATH), "images")

# Colors
C_PINK = "F4A6B3"
C_TEAL = "5B8C85"
C_LIGHT_TEAL = "7AAFA5"
C_GRAY = "A0AEC0"
C_DARK_GRAY = "718096"
C_PURPLE_BORDER = "C9B1FF"
C_LIGHT_PURPLE_BG = "FAF8FF"
C_MINT_BG = "F0FAF7"
C_CREAM_BG = "FFF8F0"
C_DIVIDER_GRAY = "E2E8F0"
C_TEXT_DARK = "4A5568"

# =============================================================================
# Helpers
# =============================================================================

def set_run_font(run, font_name="Microsoft YaHei", font_size=Pt(12), bold=False, italic=False, color=None):
    """Set font properties with eastAsia support for CJK."""
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def add_empty_para(doc, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def set_para_alignment(para, align="left"):
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_border_bottom(para, color=C_DIVIDER_GRAY, size=4, space=1):
    """Add bottom border to paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_top_bottom_border(para, color=C_DIVIDER_GRAY, size=4, space=1):
    """Add top and bottom borders."""
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_para_shading(para, color):
    """Set paragraph background shading."""
    pPr = para._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    pPr.append(shd)


def set_para_indent(para, left=Inches(0.3), right=Inches(0.3)):
    pf = para.paragraph_format
    pf.left_indent = left
    pf.right_indent = right


def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("·  ·  ·")
    set_run_font(run, font_size=Pt(14), color=C_PINK)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_body_para(doc, text, bold_phrases=None, space_after=Pt(6)):
    """Add body paragraph with optional bold highlighted phrases."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)

    if bold_phrases:
        remaining = text
        for phrase in bold_phrases:
            parts = remaining.split(phrase, 1)
            if len(parts) == 2:
                if parts[0]:
                    run = p.add_run(parts[0])
                    set_run_font(run, font_size=Pt(12))
                run = p.add_run(phrase)
                set_run_font(run, font_size=Pt(12), bold=True, color=C_PINK)
                remaining = parts[1]
        if remaining:
            run = p.add_run(remaining)
            set_run_font(run, font_size=Pt(12))
    else:
        run = p.add_run(text)
        set_run_font(run, font_size=Pt(12))
    return p


def add_quote_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    set_para_indent(p, left=Inches(0.4), right=Inches(0.4))
    set_para_shading(p, C_MINT_BG)
    add_top_bottom_border(p, color=C_DIVIDER_GRAY, size=4, space=4)
    run = p.add_run(text)
    set_run_font(run, font_size=Pt(13), color=C_DARK_GRAY, italic=True)
    return p


def add_highlight_box1(doc, text):
    """Purple border box with flower prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    set_para_indent(p, left=Inches(0.25), right=Inches(0.25))
    set_para_shading(p, C_LIGHT_PURPLE_BG)
    # Add border via table-like approach using pBdr
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="{C_PURPLE_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="{C_PURPLE_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="{C_PURPLE_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="2" w:color="{C_PURPLE_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run("✿ ")
    set_run_font(run, font_size=Pt(12), color=C_PINK)
    run = p.add_run(text)
    set_run_font(run, font_size=Pt(12), bold=True, color=C_TEXT_DARK)
    return p


def add_highlight_box2(doc, text):
    """Pink border box with cream background."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    set_para_indent(p, left=Inches(0.25), right=Inches(0.25))
    set_para_shading(p, C_CREAM_BG)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="{C_PINK}"/>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="{C_PINK}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="{C_PINK}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="2" w:color="{C_PINK}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    set_run_font(run, font_size=Pt(12), bold=True, color=C_TEAL)
    return p


def add_chapter_marker(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(str(number))
    set_run_font(run, font_size=Pt(24), bold=True, color=C_PINK)
    run = p.add_run(" / ")
    set_run_font(run, font_size=Pt(16), color=C_GRAY)
    run = p.add_run(title)
    set_run_font(run, font_size=Pt(14), bold=True, color=C_TEAL)
    # Section bottom border
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_after = Pt(8)
    border_p.paragraph_format.space_before = Pt(0)
    add_border_bottom(border_p, color=C_DIVIDER_GRAY, size=4, space=1)
    return p


def add_image_with_caption(doc, img_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.space_before = Pt(0)
    run = cap.add_run("▲ " + caption)
    set_run_font(run, font_size=Pt(10), italic=True, color=C_GRAY)
    return p


def add_end_tags(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    tags = "#职场生存 #人际关系 #自我保护 #情绪管理"
    run = p.add_run(tags)
    set_run_font(run, font_size=Pt(12), color=C_GRAY)
    return p


# =============================================================================
# Main Document Generation
# =============================================================================

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -------------------------------------------------------------------------
    # Top border (pink bottom border on first empty paragraph)
    # -------------------------------------------------------------------------
    top_p = doc.add_paragraph()
    top_p.paragraph_format.space_after = Pt(4)
    top_p.paragraph_format.space_before = Pt(0)
    add_border_bottom(top_p, color=C_PINK, size=12, space=1)

    # -------------------------------------------------------------------------
    # English label
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("WORKPLACE WISDOM")
    set_run_font(run, font_name="Georgia", font_size=Pt(10), color=C_PINK)

    # -------------------------------------------------------------------------
    # Title
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("被同事甩锅、被领导打压？")
    set_run_font(run, font_size=Pt(28), bold=True, color=C_TEAL)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_before = Pt(0)
    run = p2.add_run("学会这4招，没人再敢欺负你")
    set_run_font(run, font_size=Pt(28), bold=True, color=C_TEAL)

    # -------------------------------------------------------------------------
    # Subtitle
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("职场生存实战指南：从被动挨打到主动掌控")
    set_run_font(run, font_size=Pt(18), bold=True, color=C_LIGHT_TEAL)

    # -------------------------------------------------------------------------
    # Tagline
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("你不是不够好，只是还没学会保护自己。")
    set_run_font(run, font_size=Pt(13), italic=True, color=C_GRAY)

    # -------------------------------------------------------------------------
    # Image 1
    # -------------------------------------------------------------------------
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img1.jpg"), "职场中，学会保护自己是一门必修课")

    # -------------------------------------------------------------------------
    # Chapter 0: Opening story
    # -------------------------------------------------------------------------
    add_chapter_marker(doc, 0, "那个被甩锅的下午")

    story_text = (
        "小林入职三年，一直兢兢业业。某个周五下班前，部门经理突然在群里发飙："
        "'这个项目的客户投诉，谁负责的？'小林还没反应过来，同事阿杰就接话了："
        "'是小林跟进的，我提醒过她要注意细节。'"
        "\n\n"
        "那一刻，小林脑子嗡嗡作响。明明阿杰才是主要负责人，自己只是协助，"
        "怎么一转眼，锅就扣到了自己头上？更让她心寒的是，经理没有追问，"
        "直接@她要求周一前给出整改方案。"
        "\n\n"
        "她想说'这不是我的责任'，但话到嘴边又咽了回去。"
        "她怕被认为'推卸责任'、'没有团队精神'。"
        "于是，她默默加班整个周末，独自收拾了烂摊子。"
        "\n\n"
        "这样的场景，在职场里每天都在上演。"
        "被同事甩锅、被领导打压、被无端指责……"
        "很多人选择忍气吞声，以为'多一事不如少一事'。"
        "但真相是：你越沉默，欺负你的人越肆无忌惮。"
    )
    for para_text in story_text.split("\n\n"):
        if para_text.strip():
            add_body_para(doc, para_text.strip())

    add_divider(doc)

    # -------------------------------------------------------------------------
    # Chapter 1
    # -------------------------------------------------------------------------
    add_chapter_marker(doc, 1, "第一招：凡事留痕，让甩锅无处遁形")

    add_body_para(
        doc,
        "职场中最可怕的，不是犯错，而是'口说无凭'。"
        "当责任边界模糊时，谁有证据，谁就能掌握主动权。"
        "留痕不是心机，而是最基本的职业自保。",
        bold_phrases=["留痕不是心机，而是最基本的职业自保"]
    )

    add_highlight_box1(
        doc,
        "核心原则：所有重要沟通，尽量通过邮件、企业微信或书面形式确认，"
        "避免仅依赖口头交流。"
    )

    add_body_para(
        doc,
        "具体怎么做？首先，会议结束后立即发送会议纪要，明确每个人的分工和截止日期。"
        "其次，项目关键节点用邮件同步进展，抄送相关领导。"
        "再次，遇到口头指令，礼貌地回复：'好的，我按这个方向推进，"
        "稍后邮件确认一下细节方便吗？'"
        "最后，保存所有工作文档的版本记录，确保每一步都有据可查。"
    )

    add_body_para(
        doc,
        "有一次，我的朋友被同事在领导面前指责'拖延进度'。"
        "她不慌不忙，打开邮箱，展示了三封催促对方确认的邮件，"
        "以及对方已读未回的截图。领导当场转变了态度。"
        "记住：在职场，证据比辩解更有力量。",
        bold_phrases=["证据比辩解更有力量"]
    )

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img2.jpg"), "清晰的记录，是你最坚实的盾牌")

    # -------------------------------------------------------------------------
    # Chapter 2
    # -------------------------------------------------------------------------
    add_chapter_marker(doc, 2, "第二招：建立同盟，别做孤岛上的靶子")

    add_body_para(
        doc,
        "很多人被欺负，不是因为能力差，而是因为'势单力薄'。"
        "在职场，孤立无援的人最容易成为替罪羊。"
        "建立良好的人际网络，不是为了勾心斗角，而是为了在关键时刻有人替你说话。",
        bold_phrases=["孤立无援的人最容易成为替罪羊"]
    )

    add_highlight_box2(
        doc,
        "同盟不是拉帮结派，而是基于真诚互助的职业关系。"
        "平时多帮人一把，关键时刻才有人拉你一把。"
    )

    add_body_para(
        doc,
        "怎么建立同盟？第一，主动跨部门合作，让更多人了解你的能力和为人。"
        "第二，在同事需要帮助时伸出援手，积累人情。"
        "第三，定期与直属领导一对一沟通，汇报进展、寻求指导，"
        "让领导对你有持续的正向认知。"
        "第四，找到公司里价值观相近、能力互补的伙伴，形成互相支持的圈子。"
    )

    add_body_para(
        doc,
        "需要警惕的是，不要只依赖某一个'靠山'。"
        "职场关系瞬息万变，今天的大树可能明天就倒了。"
        "最好的策略是：让自己成为别人愿意合作的人，"
        "用专业能力和靠谱人品，构建多元化的支持网络。",
        bold_phrases=["让自己成为别人愿意合作的人"]
    )

    add_quote_block(
        doc,
        "'一个人可以走得快，但一群人才能走得远。'"
        "在职场，这句话的另一种解读是："
        "没有人撑腰的人，连站都站不稳。"
    )

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img3.jpg"), "盟友不是靠山，而是你职业路上的同行者")

    # -------------------------------------------------------------------------
    # Chapter 3
    # -------------------------------------------------------------------------
    add_chapter_marker(doc, 3, "第三招：设立边界，温柔而坚定地说'不'")

    add_body_para(
        doc,
        "很多职场霸凌，都是从'试探边界'开始的。"
        "今天让你帮忙做个表，明天让你背个锅，"
        "后天就可能把本该属于他的重大项目失误推到你身上。"
        "如果你不设立清晰的边界，别人就会不断侵犯你的领地。",
        bold_phrases=["如果你不设立清晰的边界，别人就会不断侵犯你的领地"]
    )

    add_highlight_box1(
        doc,
        "说'不'的艺术：不攻击对方，不贬低自己，只陈述事实和立场。"
    )

    add_body_para(
        doc,
        "比如同事临时甩给你一个任务，你可以说："
        "'我理解这个很急，但我手头有A项目和B项目，"
        "领导要求本周内完成。如果接这个，那两个就要延期。"
        "你看是找领导调整优先级，还是找其他人支援？'"
        "这番话既表达了拒绝，又把决策权抛回给对方，"
        "同时暗示了'我有领导交代的任务，不是无所事事'。"
    )

    add_body_para(
        doc,
        "对领导的过度要求，也可以用类似策略："
        "'我很愿意承担这个任务，但目前手头有三项工作在推进。"
        "如果接这个，我需要确认哪项可以延后，或者是否需要增加人手。'"
        "这不是推诿，而是负责任地管理预期。"
        "真正专业的领导，会欣赏这种清晰的沟通。",
        bold_phrases=["真正专业的领导，会欣赏这种清晰的沟通"]
    )

    add_body_para(
        doc,
        "设立边界还有一个重要方面：情绪边界。"
        "不要因为别人的否定就自我怀疑，不要因为领导的批评就全盘否定自己。"
        "把'事'和'人'分开：批评的是这件事的处理方式，不是你这个人的价值。"
        "守住情绪边界，才能在风暴中保持清醒。"
    )

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img4.jpg"), "温柔的边界，是对自己最大的尊重")

    # -------------------------------------------------------------------------
    # Chapter 4
    # -------------------------------------------------------------------------
    add_chapter_marker(doc, 4, "第四招：审时度势，知道何时该 escalate")

    add_body_para(
        doc,
        "前三招是防御，第四招是反击。"
        "当甩锅和打压已经常态化，当沟通无果、边界被反复践踏，"
        "你就需要考虑升级处理——向更高层反映、寻求HR介入，或者准备离开。",
        bold_phrases=["当沟通无果、边界被反复践踏，你就需要考虑升级处理"]
    )

    add_highlight_box2(
        doc,
        "Escalate 不是告状，而是基于事实的正当维权。"
        "带着证据、带着解决方案、带着冷静的情绪去沟通。"
    )

    add_body_para(
        doc,
        "升级之前，先问自己三个问题："
        "第一，我是否有充分的证据？空口无凭的投诉只会让自己显得情绪化。"
        "第二，我是否已经尝试过直接沟通？跳过直接沟通直接找上级，"
        "容易被视为'越级'或'爱打小报告'。"
        "第三，我期望的结果是什么？是调整分工、换组、还是其他解决方案？"
        "带着明确诉求去沟通，成功率会高很多。"
    )

    add_body_para(
        doc,
        "如果公司内部渠道无法解决，也要了解劳动法赋予你的权利。"
        "恶意打压、职场霸凌、违法解除劳动合同，这些都是可以依法维权的。"
        "保留好证据，咨询专业律师，必要时通过法律途径保护自己。"
        "记住：你的尊严和权益，值得被捍卫。",
        bold_phrases=["你的尊严和权益，值得被捍卫"]
    )

    add_quote_block(
        doc,
        "'不要在消耗你的地方寻找成长。'"
        "有时候，最勇敢的决定不是坚持，而是转身离开。"
    )

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img5.jpg"), "离开不是逃避，而是对自己人生的负责")

    # -------------------------------------------------------------------------
    # Mindset shift section
    # -------------------------------------------------------------------------
    add_chapter_marker(doc, 5, "心态转变：从受害者到掌控者")

    add_body_para(
        doc,
        "被欺负的人，往往有一个共同特征：内心深处觉得自己'不配被尊重'。"
        "这种信念可能来自原生家庭的打压、过往的失败经历，或者长期被PUA后的自我怀疑。"
        "但请记住：你的价值不取决于别人的评价，而取决于你对自己的认知。",
        bold_phrases=["你的价值不取决于别人的评价，而取决于你对自己的认知"]
    )

    add_body_para(
        doc,
        "第一个心态转变：从'为什么是我'到'我能做什么'。"
        "遭遇不公时，沉溺于委屈和愤怒只会消耗自己。"
        "把能量转移到解决问题上，你会发现自己比想象中更有力量。"
    )

    add_body_para(
        doc,
        "第二个心态转变：从'怕被讨厌'到'敢于被讨厌'。"
        "你不是人民币，不可能让所有人满意。"
        "当你不再把'被所有人喜欢'作为目标，"
        "你就获得了真正的自由——拒绝的自由、表达的自由、做自己的自由。",
        bold_phrases=["你就获得了真正的自由"]
    )

    add_body_para(
        doc,
        "第三个心态转变：从'等待被拯救'到'主动创造选择'。"
        "很多人忍受糟糕的职场环境，是因为'没有更好的选择'。"
        "但'没有选择'往往是一种幻觉。"
        "更新简历、拓展人脉、学习新技能、关注行业机会——"
        "当你手里有牌，你就不会害怕任何一张桌子。"
    )

    add_highlight_box1(
        doc,
        "真正的职场高手，不是从不被欺负，而是被欺负时知道如何保护自己；"
        "不是从不受伤，而是受伤后能快速恢复、变得更强大。"
    )

    # -------------------------------------------------------------------------
    # Stay vs Leave
    # -------------------------------------------------------------------------
    add_chapter_marker(doc, 6, "该留还是该走？一张决策清单")

    add_body_para(
        doc,
        "面对有毒的职场环境，很多人会陷入纠结："
        "'现在走是不是太冲动？''再忍忍会不会变好？''走了会不会找不到更好的？'"
        "以下是一些判断标准，帮助你做出理性决策。"
    )

    add_body_para(
        doc,
        "建议留下的情况："
        "第一，问题只是个别现象，整体团队和公司文化尚可。"
        "第二，你在这里还有明确的学习目标和成长空间。"
        "第三，你已经尝试用文中提到的方法改善局面，且有初步成效。"
        "第四，当前的经济状况或职业阶段，确实需要这份收入的稳定性。"
        "第五，公司有内部转岗机制，你可以尝试换到更健康的团队。"
    )

    add_body_para(
        doc,
        "建议离开的情况："
        "第一，霸凌和打压是系统性的，从领导到HR都在默许甚至参与。"
        "第二，你的身心健康已经受到明显影响，出现焦虑、失眠、抑郁等症状。"
        "第三，你已经多次尝试沟通和改进，但情况持续恶化。"
        "第四，公司业务发展明显下行，留下来只会浪费你的黄金职业期。"
        "第五，你已经有至少一个可行的下家或转型方案。",
        bold_phrases=["你的身心健康已经受到明显影响"]
    )

    add_highlight_box2(
        doc,
        "无论留下还是离开，核心原则只有一个："
        "不要让恐惧主导你的决策，要让理性、成长和对未来的期待来引领你。"
    )

    add_body_para(
        doc,
        "如果你决定留下，就把这段经历当作'职场抗压训练营'。"
        "每一次被甩锅，都是锻炼你留痕和沟通能力的机会；"
        "每一次被打压，都是检验你内心是否足够强大的试金石。"
        "如果你决定离开，就体面地告别，不要把负面情绪带到下一份工作。"
        "每一段经历都有它的价值，哪怕是教会你'什么样的环境不适合你'。"
    )

    add_divider(doc)

    # -------------------------------------------------------------------------
    # Closing
    # -------------------------------------------------------------------------
    add_body_para(
        doc,
        "职场不是童话世界，有利益的地方就有冲突，有权力的地方就有压迫。"
        "但这不意味着你要逆来顺受、任人宰割。"
        "学会留痕、建立同盟、设立边界、审时度势——"
        "这四招不是让你去算计别人，而是让你有能力保护自己。"
    )

    add_body_para(
        doc,
        "最后，想对所有正在经历职场困境的人说："
        "你不是软弱，你只是还没找到方法；"
        "你不是无能，你只是还没遇到欣赏你的舞台；"
        "你不是不够好，你只是需要学会对自己说'我值得被尊重'。"
        "从今天起，做一个不好惹的职场人——"
        "温柔有底线，善良有锋芒。",
        bold_phrases=["温柔有底线，善良有锋芒"]
    )

    # -------------------------------------------------------------------------
    # End tags
    # -------------------------------------------------------------------------
    add_end_tags(doc)

    # -------------------------------------------------------------------------
    # Save
    # -------------------------------------------------------------------------
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)

    # Copy images to output dir
    os.makedirs(OUTPUT_IMG_DIR, exist_ok=True)
    for fname in os.listdir(IMG_DIR):
        src = os.path.join(IMG_DIR, fname)
        dst = os.path.join(OUTPUT_IMG_DIR, fname)
        if os.path.isfile(src):
            shutil.copy2(src, dst)

    # Count characters
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text += para.text

    char_count = len(full_text.replace(" ", "").replace("\n", ""))
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"Total character count (no spaces): {char_count}")
    print(f"Images copied to: {OUTPUT_IMG_DIR}")


if __name__ == "__main__":
    main()
