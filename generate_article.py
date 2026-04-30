#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a 北欧杂志风 DOCX article about marriage.
"""
import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
RED = "E94560"
DARK_NAVY = "1A1A2E"
GRAY = "6B7280"
LIGHT_GRAY = "9CA3AF"
LIGHT_BORDER = "E5E7EB"
BG_LIGHT = "FAFAF5"
BG_BOX = "F0F0EA"
BG_CHAPTER = "FEF2F2"

IMG_DIR = "/tmp/article_x_imgs"
OUTPUT_DIR = "/Users/panyong/aio_project/小程序/docs/文章/2026-04-30/榆晚"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "结婚7年才明白：真正压垮婚姻的，从来不是吵架和贫穷，是因为它.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_run_font(run, font_name="宋体", size_pt=12, bold=False, italic=False, color=DARK_NAVY, eastAsia="宋体"):
    """Set font properties for a run, including eastAsia for CJK."""
    font = run.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{eastAsia}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), eastAsia)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)

def add_paragraph(doc, text="", font_name="宋体", size_pt=12, bold=False, italic=False,
                  color=DARK_NAVY, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=0, line_spacing=1.5,
                  eastAsia="宋体", first_line_indent=None,
                  top_border=False, bottom_border=False,
                  left_border=False, right_border=False,
                  border_color=RED, border_size=12,
                  bg_color=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)

    # Borders
    pPr = p._element.get_or_add_pPr()
    if top_border or bottom_border or left_border or right_border:
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                         f'  <w:top w:val="single" w:sz="{border_size if top_border else 0}" w:space="1" w:color="{border_color if top_border else "auto"}"/>'
                         f'  <w:bottom w:val="single" w:sz="{border_size if bottom_border else 0}" w:space="1" w:color="{border_color if bottom_border else "auto"}"/>'
                         f'  <w:left w:val="single" w:sz="{border_size if left_border else 0}" w:space="4" w:color="{border_color if left_border else "auto"}"/>'
                         f'  <w:right w:val="single" w:sz="{border_size if right_border else 0}" w:space="4" w:color="{border_color if right_border else "auto"}"/>'
                         f'</w:pBdr>')
        pPr.append(pBdr)

    # Background
    if bg_color:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
        pPr.append(shd)

    if text:
        run = p.add_run(text)
        set_run_font(run, font_name, size_pt, bold, italic, color, eastAsia)
    return p

def add_run_to_paragraph(p, text, font_name="宋体", size_pt=12, bold=False, italic=False,
                         color=DARK_NAVY, eastAsia="宋体"):
    run = p.add_run(text)
    set_run_font(run, font_name, size_pt, bold, italic, color, eastAsia)
    return run

def add_image_with_caption(doc, img_path, caption, width_inches=5.5):
    """Add left-aligned image with italic caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    if os.path.exists(img_path):
        run.add_picture(img_path, width=Inches(width_inches))
    # Caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(12)
    cr = cp.add_run(f"▲ {caption}")
    set_run_font(cr, "Georgia", 10, False, True, LIGHT_GRAY, "Georgia")

def add_center_dot(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("·")
    set_run_font(r, "Georgia", 18, False, False, RED, "Georgia")

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("━━━━")
    set_run_font(r, "Georgia", 10, False, False, RED, "Georgia")

def add_quote_block(doc, text):
    p = add_paragraph(doc, text, "宋体", 13, False, False, GRAY,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=12, space_after=12, line_spacing=1.6,
                      left_border=True, border_color=LIGHT_BORDER, border_size=18,
                      bg_color=BG_LIGHT)
    return p

def add_highlight_box(doc, text):
    p = add_paragraph(doc, "", "宋体", 14, False, False, DARK_NAVY,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=12, space_after=12, line_spacing=1.6,
                      bg_color=BG_BOX)
    r1 = p.add_run("· ")
    set_run_font(r1, "宋体", 14, True, False, RED, "宋体")
    r2 = p.add_run(text)
    set_run_font(r2, "宋体", 14, False, False, DARK_NAVY, "宋体")
    return p

def add_chapter_highlight(doc, text):
    p = add_paragraph(doc, text, "宋体", 13, True, False, DARK_NAVY,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=10, space_after=10, line_spacing=1.6,
                      left_border=True, border_color=RED, border_size=24,
                      bg_color=BG_CHAPTER)
    return p

def add_chapter_marker(doc, number, title):
    # Number
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(number)
    set_run_font(r, "Georgia", 28, True, False, RED, "Georgia")
    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(8)
    tr = tp.add_run(title)
    set_run_font(tr, "宋体", 15, True, False, DARK_NAVY, "宋体")

def add_center_quote(doc, quote, attribution):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"\"{quote}\"")
    set_run_font(r, "Georgia", 13, False, True, GRAY, "Georgia")
    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after = Pt(16)
    ar = ap.add_run(f"— {attribution}")
    set_run_font(ar, "Georgia", 11, False, False, LIGHT_GRAY, "Georgia")

def add_body_paragraph(doc, text, highlight_phrases=None):
    """Add body text with optional red highlights."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.first_line_indent = Cm(0.74)

    if highlight_phrases:
        remaining = text
        for phrase in highlight_phrases:
            parts = remaining.split(phrase, 1)
            if len(parts) == 2:
                if parts[0]:
                    r = p.add_run(parts[0])
                    set_run_font(r, "宋体", 12, False, False, DARK_NAVY, "宋体")
                r = p.add_run(phrase)
                set_run_font(r, "宋体", 12, False, False, RED, "宋体")
                remaining = parts[1]
        if remaining:
            r = p.add_run(remaining)
            set_run_font(r, "宋体", 12, False, False, DARK_NAVY, "宋体")
    else:
        r = p.add_run(text)
        set_run_font(r, "宋体", 12, False, False, DARK_NAVY, "宋体")
    return p

# ---------------------------------------------------------------------------
# Main document generation
# ---------------------------------------------------------------------------
def main():
    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== English label =====
    add_paragraph(doc, "MARRIAGE & EMOTION", "Georgia", 10, False, False, LIGHT_GRAY,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6)

    # ===== Title =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(8)
    # Top border
    pPr = title_p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                     f'  <w:top w:val="single" w:sz="24" w:space="4" w:color="{RED}"/>'
                     f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                     f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                     f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                     f'</w:pBdr>')
    pPr.append(pBdr)
    tr = title_p.add_run("结婚7年才明白：真正压垮婚姻的，从来不是吵架和贫穷，是因为它")
    set_run_font(tr, "宋体", 28, True, False, DARK_NAVY, "宋体")

    # ===== Divider =====
    add_divider(doc)

    # ===== Subtitle =====
    add_paragraph(doc, "情感忽视：婚姻中最隐蔽的杀手", "宋体", 16, False, False, RED,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6)

    # ===== Tagline =====
    add_paragraph(doc, "当沉默取代了争吵，当冷漠吞噬了热情，婚姻便在无声中走向终结。", "宋体", 13, False, True, GRAY,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=12)

    # ===== Image 1 =====
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img1.jpg"),
                           "婚姻中的沉默，往往比争吵更令人窒息", 5.5)

    # ===== Opening =====
    add_body_paragraph(doc,
        "我和丈夫结婚七年，在外人眼里，我们是令人羡慕的一对。没有激烈的争吵，没有经济的困顿，"
        "房子车子都有，孩子乖巧懂事。可只有我自己知道，这段婚姻正在以一种看不见的方式，"
        "一点一点地腐烂、崩塌。",
        ["一点一点地腐烂、崩塌"])

    add_body_paragraph(doc,
        "曾经我以为，婚姻最大的敌人是贫穷和争吵。小时候看父母为了钱吵得不可开交，"
        "我便暗暗发誓，以后一定要找一个经济条件不错、性格温和的人。"
        "我以为只要避开这两样，婚姻就能安稳长久。可七年的婚姻生活告诉我："
        "真正压垮婚姻的，从来不是吵架和贫穷，而是情感忽视。",
        ["真正压垮婚姻的，从来不是吵架和贫穷，而是情感忽视"])

    add_center_dot(doc)

    # ===== Chapter 01 =====
    add_chapter_marker(doc, "01", "那些看似平静的日子，其实是感情流失的过程")

    add_body_paragraph(doc,
        "结婚第一年，我们还会为了今天吃什么、周末去哪里玩而讨论半天。"
        "他会记得我不爱吃香菜，我会记得他喜欢喝冰美式。"
        "那时候，我们之间有说不完的话，哪怕是废话，也说得津津有味。")

    add_body_paragraph(doc,
        "可不知从什么时候开始，这些细碎的交流渐渐消失了。"
        "他下班回家，径直走向沙发，打开手机刷短视频。"
        "我在厨房忙碌，喊他帮忙递个东西，他应了一声，却迟迟没有动静。"
        "等我端着菜出来，他依旧盯着屏幕，仿佛我根本不存在。",
        ["仿佛我根本不存在"])

    add_body_paragraph(doc,
        "我开始学会不喊他，自己做完所有的事。"
        "我开始学会不问他今天过得怎么样，因为问了也是敷衍的\"还行\"。"
        "我开始学会不跟他分享工作中的趣事，因为他头也不抬地\"嗯\"一声，"
        "那种被漠视的感觉，比直接拒绝更让人心寒。")

    add_quote_block(doc,
        "\"婚姻中最可怕的不是争吵，而是连争吵的欲望都没有了。\""
        "当两个人在同一屋檐下，却像两个陌生人一样各自生活，"
        "那种孤独感，比单身时更甚。")

    add_body_paragraph(doc,
        "我试过跟他沟通，可他总觉得我在小题大做。"
        "\"我又没出轨，又没家暴，你还有什么不满意的？\""
        "这句话像一盆冷水，浇灭了我所有想要改变的念头。"
        "是啊，他没有做错什么，可他也什么都没有做对。"
        "他没有看见我的疲惫，没有听见我的需求，"
        "没有在我需要的时候，给过我一个拥抱。")

    # ===== Image 2 =====
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img2.jpg"),
                           "两个人在同一屋檐下，却像陌生人一样生活", 5.5)

    add_center_dot(doc)

    # ===== Chapter 02 =====
    add_chapter_marker(doc, "02", "为什么吵架和贫穷，反而不是婚姻的杀手")

    add_body_paragraph(doc,
        "很多人不理解，为什么我说吵架和贫穷不是婚姻最大的敌人。"
        "其实，吵架至少说明双方还在乎，还在试图让对方理解自己。"
        "贫穷至少让两个人有了共同奋斗的目标，有了相依为命的紧密感。")

    add_body_paragraph(doc,
        "我有一对朋友，夫妻俩经常为了鸡毛蒜皮的小事吵得面红耳赤。"
        "可吵完之后，他们会坐下来好好谈，会道歉，会拥抱，会和好如初。"
        "他们的婚姻虽然不够\"体面\"，但感情却在一次次争吵中变得更加真实和牢固。")

    add_highlight_box(doc, "吵架是沟通的极端形式，而冷漠是沟通的彻底终结。")

    add_body_paragraph(doc,
        "再说贫穷。我父母那一代人，很多人都是在贫困中携手走过来的。"
        "他们没有钱买礼物，但父亲会在母亲生日那天早起煮一碗红糖鸡蛋；"
        "他们买不起新衣，但母亲会把父亲磨破的衬衫一针一线缝补好。"
        "那种在困顿中相互扶持的温暖，是任何物质都无法替代的。")

    add_body_paragraph(doc,
        "反观我们这些看似条件不错的夫妻，"
        "有了独立的卧室，有了各自的手机，有了不需要对方也能过得很好的能力，"
        "却失去了最珍贵的东西——心与心的连接。"
        "我们在物质上越来越富足，在情感上却越来越贫瘠。",
        ["心与心的连接"])

    add_center_dot(doc)

    # ===== Chapter 03 =====
    add_chapter_marker(doc, "03", "情感忽视的五个日常信号")

    add_chapter_highlight(doc, "如果你发现婚姻中出现以下迹象，请警惕：情感忽视正在侵蚀你们的感情。")

    signals = [
        ("回应敷衍", "你兴致勃勃地分享一件事，对方却只是\"嗯\"\"哦\"地应付，"
         "眼神始终停留在手机或电视上，让你觉得自己在说废话。"),
        ("回避沟通", "每当你试图谈论关系中的问题，对方要么沉默不语，要么转移话题，"
         "要么直接说\"你想多了\"，让你感到自己的感受不被重视。"),
        ("缺乏肢体接触", "曾经的拥抱、牵手、亲吻渐渐消失，"
         "即使在同一张床上，也像是隔着一道无形的墙。"),
        ("忽视重要日子", "生日、纪念日、情人节，不再有惊喜，甚至不再被记得。"
         "你提醒了，对方才恍然大悟；你不说，就当作什么都没发生。"),
        ("情感支持缺位", "你遇到困难或委屈时，对方不会主动关心，"
         "更不会站在你身边。你感觉自己是在独自面对这个世界，"
         "而那个本该与你并肩的人，却选择了袖手旁观。")
    ]

    for title, desc in signals:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tp.paragraph_format.space_before = Pt(10)
        tp.paragraph_format.space_after = Pt(2)
        tp.paragraph_format.line_spacing = 1.6
        tr = tp.add_run(f"· {title}")
        set_run_font(tr, "宋体", 13, True, False, RED, "宋体")
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dp.paragraph_format.space_before = Pt(0)
        dp.paragraph_format.space_after = Pt(6)
        dp.paragraph_format.line_spacing = 1.6
        dp.paragraph_format.first_line_indent = Cm(0.74)
        dr = dp.add_run(desc)
        set_run_font(dr, "宋体", 12, False, False, DARK_NAVY, "宋体")

    # ===== Image 3 =====
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img3.jpg"),
                           "忽视，是婚姻中最隐蔽却最致命的伤口", 5.5)

    add_center_dot(doc)

    # ===== Chapter 04 =====
    add_chapter_marker(doc, "04", "我是如何意识到问题，并试图改变的")

    add_body_paragraph(doc,
        "改变发生在结婚第六年的一个冬夜。"
        "那天我加班到很晚，回家的路上下起了大雪。"
        "我骑着电动车，寒风刺骨，手冻得几乎握不住车把。"
        "我给丈夫打电话，想让他开车来接我，可电话那头传来的是游戏的声音，"
        "他含糊地说\"你自己打车回来吧\"，然后匆匆挂断了。")

    add_body_paragraph(doc,
        "那一刻，我站在路灯下，雪花落在脸上，凉得刺骨。"
        "我突然意识到，在这段婚姻里，我一直都是一个人。"
        "一个人做饭，一个人带孩子，一个人面对生活的风雨。"
        "他就在那里，却像空气一样，既无处不在，又毫无存在感。",
        ["既无处不在，又毫无存在感"])

    add_center_quote(doc,
        "世界上最遥远的距离，不是生与死，而是我就站在你面前，你却看不见我。",
        "改编自泰戈尔")

    add_body_paragraph(doc,
        "回家后，我写了一封长信给他。"
        "我没有指责，只是平静地讲述这些年我的感受——"
        "那些被忽视的瞬间，那些渴望回应却没有得到的日子，"
        "那种明明有伴侣却活得像单亲妈妈的孤独。")

    add_body_paragraph(doc,
        "他读完信，沉默了很久。"
        "然后他说了一句让我至今难忘的话："
        "\"我以为只要赚钱养家，就是对你好了。我没想到，你需要的不是钱，是我。\""
        "那一刻，我看到了他眼中的愧疚，也看到了一丝改变的希望。")

    add_highlight_box(doc, "很多时候，情感忽视并非出于恶意，而是源于无知。"
                          "对方可能根本不知道你需要什么，因为你们从未真正谈论过。")

    # ===== Image 4 =====
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img4.jpg"),
                           "改变，从一次真诚的对话开始", 5.5)

    add_center_dot(doc)

    # ===== Chapter 05 =====
    add_chapter_marker(doc, "05", "重建连接：给婚姻中每一对夫妻的建议")

    add_body_paragraph(doc,
        "这一年，我们在努力修复这段关系。"
        "过程并不容易，旧有的模式像惯性一样难以打破。"
        "但我们都在尝试，都在学习如何重新看见对方。")

    add_chapter_highlight(doc, "以下是我们实践过的一些方法，希望对你们也有帮助：")

    tips = [
        ("每天15分钟的高质量对话", "放下手机，看着对方的眼睛，"
         "聊聊今天发生了什么，有什么感受。不是汇报工作，而是分享心情。"),
        ("重建肢体接触", "从牵手开始，从拥抱开始。"
         "肢体接触是情感连接最原始也最直接的方式，不要小看一个拥抱的力量。"),
        ("表达感谢和欣赏", "即使是小事，也要说出来。"
         "\"谢谢你今天洗碗\"\"我喜欢你穿这件衣服的样子\"，"
         "这些话语像润滑剂，让关系保持柔软和温暖。"),
        ("共同做一件事", "一起做饭、一起散步、一起看电影。"
         "共同的经历会创造新的记忆，也会让你们重新找到\"我们\"的感觉。"),
        ("定期检视关系", "每个月找一个安静的夜晚，"
         "坦诚地聊聊彼此的感受和需求，及时调整，不让问题积累。")
    ]

    for title, desc in tips:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tp.paragraph_format.space_before = Pt(10)
        tp.paragraph_format.space_after = Pt(2)
        tp.paragraph_format.line_spacing = 1.6
        tr = tp.add_run(f"· {title}")
        set_run_font(tr, "宋体", 13, True, False, RED, "宋体")
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dp.paragraph_format.space_before = Pt(0)
        dp.paragraph_format.space_after = Pt(6)
        dp.paragraph_format.line_spacing = 1.6
        dp.paragraph_format.first_line_indent = Cm(0.74)
        dr = dp.add_run(desc)
        set_run_font(dr, "宋体", 12, False, False, DARK_NAVY, "宋体")

    # ===== Image 5 =====
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img5.jpg"),
                           "重建连接，是一场需要两个人共同努力的旅程", 5.5)

    add_center_dot(doc)

    # ===== Closing =====
    add_body_paragraph(doc,
        "结婚七年，我终于明白："
        "婚姻不是一场一劳永逸的交易，而是一段需要持续经营的关系。"
        "贫穷可以一起克服，争吵可以一起化解，"
        "但情感忽视却像慢性毒药，在不知不觉中侵蚀着爱的根基。")

    add_body_paragraph(doc,
        "如果你也正在经历这样的婚姻，请不要绝望。"
        "改变永远不晚，只要你们还愿意看见对方，还愿意为了这段关系努力。"
        "因为婚姻最珍贵的，不是那张结婚证，"
        "而是两个灵魂在漫长岁月中，依然选择彼此靠近的勇气和决心。",
        ["两个灵魂在漫长岁月中，依然选择彼此靠近的勇气和决心"])

    add_quote_block(doc,
        "\"爱不是寻找一个完美的人，而是学会用完美的眼光，欣赏那个并不完美的人。\""
        "愿每一对夫妻，都能在柴米油盐的日常中，不忘当初相爱的初心。")

    # ===== Bottom border section =====
    bp = doc.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bp.paragraph_format.space_before = Pt(24)
    bp.paragraph_format.space_after = Pt(12)
    pPr = bp._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                     f'  <w:bottom w:val="single" w:sz="18" w:space="4" w:color="{RED}"/>'
                     f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                     f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                     f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                     f'</w:pBdr>')
    pPr.append(pBdr)
    br = bp.add_run("END")
    set_run_font(br, "Georgia", 12, True, False, RED, "Georgia")

    # ===== End tags =====
    add_paragraph(doc, "#婚姻感悟 #情感忽视 #夫妻关系 #七年之痒 #婚姻修复",
                  "宋体", 11, False, False, LIGHT_GRAY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"DOCX saved to: {OUTPUT_FILE}")

    # Copy images to output dir
    img_out = os.path.join(OUTPUT_DIR, "images")
    os.makedirs(img_out, exist_ok=True)
    for i in range(1, 6):
        src = os.path.join(IMG_DIR, f"img{i}.jpg")
        if os.path.exists(src):
            shutil.copy2(src, os.path.join(img_out, f"img{i}.jpg"))
    print(f"Images copied to: {img_out}")

    # Character count
    full_text = ""
    for p in doc.paragraphs:
        full_text += p.text
    print(f"Total character count (approx): {len(full_text)}")

if __name__ == "__main__":
    main()
