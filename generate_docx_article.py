#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a docx article in 工业水泥风 (Industrial Cement) style.
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# Style Constants
# ============================================================
GOLD = RGBColor(0xFF, 0xD7, 0x00)
DARK_GRAY = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x8A, 0x8A, 0x8A)
MEDIUM_GRAY = RGBColor(0x5A, 0x5A, 0x5A)
LIGHT_GRAY_BG = "EBEBEB"
CREAM_BG = "FFF8DC"
SILVER = RGBColor(0xC0, 0xC0, 0xC0)
NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)

IMG_DIR = "/tmp/article_r_imgs"
OUT_DIR = "/Users/panyong/aio_project/小程序/docs/文章/2026-04-30/xieyh"
DOCX_PATH = os.path.join(OUT_DIR, "面试时这句话一说，HR立刻给你涨薪30%.docx")


def set_run_font(run, font_name="Microsoft YaHei", size_pt=None, bold=False, italic=False, color=None, eastAsia=True):
    """Set font properties for a run, with eastAsia support for CJK."""
    font = run.font
    if size_pt:
        font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
    run.font.name = font_name
    if eastAsia:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), font_name)


def add_label(doc, text="[ EFFICIENCY LAB ]"):
    """Add the label paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=10, color=GRAY)
    p.space_after = Pt(6)
    return p


def add_title(doc, text):
    """Add main title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=28, bold=True, color=DARK_GRAY)
    p.space_after = Pt(12)
    return p


def add_subtitle(doc, text):
    """Add subtitle in gold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=22, bold=True, color=GOLD)
    p.space_after = Pt(10)
    return p


def add_tagline(doc, text):
    """Add tagline."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=13, italic=True, color=MEDIUM_GRAY)
    p.space_after = Pt(16)
    return p


def add_divider(doc):
    """Add centered divider: ——  ●  ——"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("——  ●  ——")
    set_run_font(run, size_pt=14, color=GOLD)
    p.space_before = Pt(12)
    p.space_after = Pt(12)
    return p


def add_image_with_caption(doc, img_path, caption):
    """Add centered image with light gray background and caption."""
    # We simulate background by adding a table with gray shading
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    # Set cell shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_GRAY_BG}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    # Add image
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    # Caption below table
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    set_run_font(cap_run, size_pt=10, italic=True, color=GRAY)
    cap_p.space_after = Pt(14)
    return table


def add_quote_block(doc, text):
    """Quote block: left border gold, background #EBEBEB, 13pt medium gray."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    # Background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_GRAY_BG}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    # Left border gold
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="FFD700"/>'
        f'  <w:top w:val="nil"/>'
        f'  <w:bottom w:val="nil"/>'
        f'  <w:right w:val="nil"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=13, color=MEDIUM_GRAY)
    p.space_after = Pt(8)
    doc.add_paragraph().space_after = Pt(6)
    return table


def add_highlight_box(doc, text):
    """Highlight box: left border gold, background #FFF8DC, 12pt bold dark gray."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CREAM_BG}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="FFD700"/>'
        f'  <w:top w:val="nil"/>'
        f'  <w:bottom w:val="nil"/>'
        f'  <w:right w:val="nil"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=12, bold=True, color=DARK_GRAY)
    p.space_after = Pt(8)
    doc.add_paragraph().space_after = Pt(6)
    return table


def add_chapter_marker(doc, number_text):
    """Chapter marker like [ 01 ] in gold, number 24pt, brackets 14pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_bracket1 = p.add_run("[ ")
    set_run_font(run_bracket1, size_pt=14, color=GOLD)
    run_num = p.add_run(number_text)
    set_run_font(run_num, size_pt=24, bold=True, color=GOLD)
    run_bracket2 = p.add_run(" ]")
    set_run_font(run_bracket2, size_pt=14, color=GOLD)
    p.space_after = Pt(4)
    return p


def add_chapter_title(doc, text):
    """Chapter title: 15pt bold dark gray."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=15, bold=True, color=DARK_GRAY)
    p.space_after = Pt(10)
    return p


def add_body_paragraph(doc, text, bold=False, color=None, size_pt=12):
    """Standard body text paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold, color=color if color else DARK_GRAY)
    p.space_after = Pt(10)
    return p


def add_bullet(doc, text):
    """Bullet point: ▸ in gold bold, then text bold near-black."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run_bullet = p.add_run("▸ ")
    set_run_font(run_bullet, size_pt=12, bold=True, color=GOLD)
    run_text = p.add_run(text)
    set_run_font(run_text, size_pt=12, bold=True, color=NEAR_BLACK)
    p.space_after = Pt(8)
    return p


def add_section_bottom_border(doc):
    """Add a thin silver bottom border to a paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="C0C0C0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    p.space_after = Pt(12)
    return p


def add_end_tags(doc, tags):
    """End tags centered, e.g. #面试技巧 #薪资谈判 #职场成长"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, tag in enumerate(tags):
        if i > 0:
            p.add_run("  ")
        run = p.add_run(tag)
        set_run_font(run, size_pt=12, color=MEDIUM_GRAY)
    p.space_before = Pt(20)
    p.space_after = Pt(10)
    return p


# ============================================================
# Document Assembly
# ============================================================
def build_document():
    doc = Document()

    # Page margins
    sections = doc.sections[0]
    sections.top_margin = Cm(2.5)
    sections.bottom_margin = Cm(2.5)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    # --- Header / Label ---
    add_label(doc, "[ EFFICIENCY LAB ]")

    # --- Title ---
    add_title(doc, "面试时这句话一说，HR立刻给你涨薪30%")

    # --- Subtitle ---
    add_subtitle(doc, "薪资谈判不是博弈，而是价值呈现的艺术")

    # --- Tagline ---
    add_tagline(doc, "掌握核心话术，让你的下一份工作起点更高")

    add_divider(doc)

    # --- Image 1 ---
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img1_interview.jpg"),
                           "每一次面试，都是一次重新定义自己价值的机会")

    # --- Chapter 01 ---
    add_chapter_marker(doc, "01")
    add_chapter_title(doc, '一个真实的故事：从"不敢谈"到"多拿30%"')
    add_body_paragraph(doc,
        "去年秋天，我的朋友小林参加了一家互联网大厂的面试。"
        "经过三轮技术面试，他顺利进入了HR面。"
        "HR开出的薪资是年薪24万，比他上一份工作的18万已经高出了33%。"
        "但小林心里清楚，以他的技术能力和项目经验，这个报价其实偏低。"
        "他犹豫了很久，最终鼓起勇气说出了那句话。"
        "结果HR沉默了三秒，然后说：'我去跟领导申请一下。'"
        "十分钟后，HR回来了，新offer变成了年薪31万——涨幅超过70%。")
    add_quote_block(doc,
        "\"真正限制你薪资的，从来不是市场行情，而是你不敢开口的勇气。\"")
    add_body_paragraph(doc,
        "小林后来告诉我，那一刻他明白了："
        "薪资谈判不是乞讨，也不是对抗，而是一次平等的价值对话。"
        "当你能够清晰、自信地表达自己的价值时，HR不仅不会反感，反而会尊重你。")
    add_section_bottom_border(doc)

    # --- Chapter 02 ---
    add_chapter_marker(doc, "02")
    add_chapter_title(doc, "薪资谈判的三大核心原则")
    add_body_paragraph(doc,
        "在进入具体话术之前，你需要先理解薪资谈判的底层逻辑。"
        "很多人谈薪失败，不是因为不会说话，而是因为搞错了谈判的本质。")
    add_highlight_box(doc,
        "原则一：先谈价值，再谈数字。HR关心的不是你想要什么，而是你值多少。")
    add_body_paragraph(doc,
        "在提出任何薪资要求之前，你必须先让对方认可你的价值。"
        "这意味着你需要在面试过程中，通过具体的项目案例、数据成果来建立你的'价值锚点'。"
        "比如：'我在上一家公司主导的用户增长项目，三个月内将DAU提升了40%。'")
    add_highlight_box(doc,
        "原则二：永远不要让对方先知道你目前的薪资。这是谈判中最致命的弱点。")
    add_body_paragraph(doc,
        "当HR问'你目前的薪资是多少'时，很多人的第一反应是如实回答。"
        "但这等于提前亮出了底牌。正确的做法是转移焦点："
        "'我更看重这个岗位的发展空间，相信贵司有合理的薪酬体系。'")
    add_highlight_box(doc,
        "原则三：给出一个范围，而不是一个具体数字。范围的上限要略高于你的真实期望。")
    add_body_paragraph(doc,
        "比如你的期望是30万，你可以说：'基于我的经验和市场行情，我的期望在32万到38万之间。'")
    add_section_bottom_border(doc)

    # --- Image 2 ---
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img2_business.jpg"),
                           "谈判桌上，准备充分的人永远占据主动权")

    # --- Chapter 03 ---
    add_chapter_marker(doc, "03")
    add_chapter_title(doc, '那句让HR立刻涨薪的"魔法话术"')
    add_body_paragraph(doc,
        "好了，现在来到本文最核心的部分——那句话到底是什么？")
    add_highlight_box(doc,
        "\"我非常看好这个岗位和团队，也相信我的能力能够为贵司创造显著价值。"
        "基于我对这个岗位的理解以及目前的市场行情，我希望年薪能够达到XX万。"
        "当然，我也理解薪酬是一个综合体系，如果这个数字有挑战，"
        "我想了解一下贵司在绩效奖金、股权激励或其他福利方面是否有补充空间？\"")
    add_body_paragraph(doc,
        "这句话的精妙之处在于，它同时完成了三个任务：")
    add_bullet(doc, "表达了强烈的入职意愿，让HR感受到你的诚意")
    add_bullet(doc, "明确提出了具体的薪资期望，不给对方模糊的空间")
    add_bullet(doc, "展示了灵活性，为后续的谈判留出了回旋余地")
    add_body_paragraph(doc,
        "HR听到这句话后，通常会经历一个心理过程："
        "首先，你的积极态度降低了对方的防御心理；"
        "其次，你给出的具体数字让对方有了明确的参照；"
        "最后，你提到的'补充空间'给了HR一个台阶——"
        "即使基础薪资达不到，也可以通过其他方式满足你。")
    add_quote_block(doc,
        "\"最高明的谈判，不是让对方觉得输了，而是让对方觉得找到了双赢的方案。\"")
    add_section_bottom_border(doc)

    # --- Image 3 ---
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img3_negotiation.jpg"),
                           "每一次开口，都是一次精心设计的价值传递")

    # --- Chapter 04 ---
    add_chapter_marker(doc, "04")
    add_chapter_title(doc, "谈薪时最常见的五个致命错误")
    add_body_paragraph(doc,
        "知道了该说什么，同样重要的是知道不该说什么。"
        "以下是我在数百个案例中总结出的五个最常见错误：")
    add_bullet(doc, "错误一：过早谈薪。在HR还没认可你的价值之前就抛出数字，等于自断后路。")
    add_bullet(doc, "错误二：只谈钱不谈价值。如果你给不出理由，任何数字都显得贪婪。")
    add_bullet(doc, "错误三：威胁式谈判。'如果不给XX万我就去别家'——这是最愚蠢的策略。")
    add_bullet(doc, "错误四：接受口头offer。所有承诺都必须落到书面，否则等于没有。")
    add_bullet(doc, "错误五：忽视总包。基础薪资只是总包的一部分，股票、期权、奖金同样重要。")
    add_body_paragraph(doc,
        "很多人谈薪失败，不是因为能力不够，而是因为犯了这些本可以避免的错误。"
        "记住：每一次谈判都是一场信息战，准备越充分，胜算越大。")
    add_section_bottom_border(doc)

    # --- Image 4 ---
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img4_handshake.jpg"),
                           "握手的那一刻，是价值被认可的最好证明")

    # --- Chapter 05 ---
    add_chapter_marker(doc, "05")
    add_chapter_title(doc, "实战锦囊：让你谈判成功率翻倍的具体技巧")
    add_body_paragraph(doc,
        "理论讲完了，接下来是一些可以直接套用的实战技巧。"
        "把这些技巧融入你的谈判流程，你的成功率至少能提升一倍。")
    add_highlight_box(doc,
        "技巧一：提前做市场调研。了解目标岗位在行业内的薪资中位数和高位数。")
    add_body_paragraph(doc,
        "你可以通过招聘网站、行业报告、脉脉、LinkedIn Salary等渠道获取数据。"
        "谈判时引用这些数据，会让你的要求显得有理有据，而不是漫天要价。")
    add_highlight_box(doc,
        "技巧二：准备三个数字：理想值、可接受值、底线值。")
    add_body_paragraph(doc,
        "理想值是你最希望拿到的数字，可接受值是你觉得合理的数字，底线值是你绝对不能突破的数字。"
        "谈判时，从理想值开始，逐步向可接受值靠拢，但绝不触碰底线。")
    add_highlight_box(doc,
        '技巧三：用"我们"代替"我"，把谈判变成合作而不是对抗。')
    add_body_paragraph(doc,
        "比如：'我们来看看怎样才能达成一个双方都满意的结果。'"
        "这种措辞能够软化对方的立场，让对话更加顺畅。")
    add_highlight_box(doc,
        "技巧四：学会沉默。提出数字后，不要急着解释，给对方思考的时间。")
    add_body_paragraph(doc,
        "沉默是一种强大的谈判武器。当你说完期望薪资后，安静地看着对方。"
        "谁先开口，谁就失去了主动权。")
    add_highlight_box(doc,
        "技巧五：准备Plan B。即使谈判失败，你也要有其他选择。")
    add_body_paragraph(doc,
        "最好的谈判筹码，是你根本不在乎这次谈判的结果。"
        "当你手上有多个offer时，你的底气会完全不同。")
    add_section_bottom_border(doc)

    # --- Image 5 ---
    add_image_with_caption(doc, os.path.join(IMG_DIR, "img5_success.jpg"),
                           "成功的谈判，始于充分的准备和坚定的自信")

    # --- Chapter 06 (Bonus) ---
    add_chapter_marker(doc, "06")
    add_chapter_title(doc, "写在最后：你的价值，值得被认真对待")
    add_body_paragraph(doc,
        "薪资谈判的本质，不是一场零和博弈，而是一次价值对齐的过程。"
        "当你能够清晰、自信、有策略地表达自己的价值时，"
        "你不仅是在争取更高的薪水，更是在建立一种平等、专业的职场关系。")
    add_quote_block(doc,
        "\"你开口要的薪资，决定了别人如何看待你的价值。"
        "低估自己，是对自己最大的不公平。\"")
    add_body_paragraph(doc,
        "希望这篇文章能够帮助你在下一次面试中，"
        "勇敢地说出那句话，拿到你应得的回报。"
        "记住：市场永远愿意为真正的价值买单，"
        "关键在于，你是否准备好了去争取。"
        "很多人之所以在薪资谈判中处于被动，根本原因在于缺乏系统性的准备。"
        "他们临时抱佛脚，在HR面前语无伦次，要么不敢开口，要么开口就错。"
        "真正的高手，在面试前就已经完成了所有的功课："
        "市场调研、价值梳理、话术演练、心理建设。"
        "当机会来临时，他们只需要从容地展示自己，然后优雅地提出期望。"
        "这不是天赋，而是可以学习和训练的技能。"
        "从今天开始，认真对待每一次面试，认真对待自己的价值。"
        "你的职业生涯，值得一个更高的起点。")
    add_highlight_box(doc,
        "行动清单：下次面试前，写下你的三个数字，练习三遍话术，然后自信地去谈。"
        "记住：敢于开口的人，已经赢了一半。")

    add_divider(doc)

    # --- End tags ---
    add_end_tags(doc, ["#面试技巧", "#薪资谈判", "#职场成长", "#升职加薪"])

    # Save
    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(DOCX_PATH)
    return DOCX_PATH


def count_characters_in_docx(path):
    """Rough character count by extracting all paragraph text."""
    doc = Document(path)
    total = 0
    for para in doc.paragraphs:
        total += len(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += len(para.text)
    return total


if __name__ == "__main__":
    build_document()
    char_count = count_characters_in_docx(DOCX_PATH)
    print(f"DOCX saved to: {DOCX_PATH}")
    print(f"Total characters: {char_count}")
    # Copy images to output directory
    img_out = os.path.join(OUT_DIR, "images")
    os.makedirs(img_out, exist_ok=True)
    for f in os.listdir(IMG_DIR):
        shutil.copy2(os.path.join(IMG_DIR, f), os.path.join(img_out, f))
    print(f"Images copied to: {img_out}")
