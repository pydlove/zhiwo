#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a styled DOCX article for the given title in 基础风格.
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Config ──────────────────────────────────────────────────────────────
OUTPUT_DIR = "/Users/panyong/aio_project/小程序/docs/文章/2026-04-30/颂頌"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "年薪30万后才明白的职场真相：能力不是最重要的，重要的是它.docx")
IMG_DIR = "/tmp/article_y_imgs"

# Colors
RED = RGBColor(0xE7, 0x4C, 0x3C)
DARK_BLUE = RGBColor(0x2C, 0x3E, 0x50)
GRAY = RGBColor(0x95, 0xA5, 0xA6)
BLUE_GRAY = RGBColor(0x5D, 0x6D, 0x7E)
LIGHT_BG = RGBColor(0xF8, 0xF9, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set run font ────────────────────────────────────────────────
def set_run_font(run, font_name="Microsoft YaHei", size_pt=None, bold=False, color=None, italic=False):
    font = run.font
    if size_pt:
        font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# ── Helper: add paragraph ───────────────────────────────────────────────
def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False, color=None,
             italic=False, space_before=0, space_after=0, line_spacing=1.5, font_name="Microsoft YaHei"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, size_pt=size, bold=bold, color=color, italic=italic)
    return p

# ── Helper: add divider ─────────────────────────────────────────────────
def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("—— ◆ ——")
    set_run_font(run, size_pt=11, bold=False, color=RED)
    return p

# ── Helper: add quote block ─────────────────────────────────────────────
def add_quote_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    # Left border via shading/paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="12" w:space="4" w:color="E74C3C"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    # Background shading
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F8F9FA"/>'
    )
    pPr.append(shd)
    run = p.add_run(text)
    set_run_font(run, size_pt=13, color=BLUE_GRAY)
    return p

# ── Helper: add highlight box ───────────────────────────────────────────
def add_highlight_box(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="1" w:color="ECF0F1"/>'
        '  <w:left w:val="single" w:sz="12" w:space="4" w:color="E74C3C"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="ECF0F1"/>'
        '  <w:right w:val="single" w:sz="4" w:space="1" w:color="ECF0F1"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F8F9FA"/>'
    )
    pPr.append(shd)
    run1 = p.add_run("✦ ")
    set_run_font(run1, size_pt=12, color=RED, bold=True)
    run2 = p.add_run(text)
    set_run_font(run2, size_pt=12, color=DARK_BLUE)
    return p

# ── Helper: add image with caption ──────────────────────────────────────
def add_image_with_caption(doc, img_path, caption_text, width_inches=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(10)
    run_cap = cap_p.add_run(f"▲ {caption_text}")
    set_run_font(run_cap, size_pt=10, italic=True, color=GRAY)
    return p

# ── Helper: add body paragraph with bold highlight ──────────────────────
def add_body_with_highlight(doc, segments):
    """
    segments: list of tuples (text, is_highlight)
    is_highlight means bold + RED color
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars indent
    for text, is_highlight in segments:
        run = p.add_run(text)
        if is_highlight:
            set_run_font(run, size_pt=12, bold=True, color=RED)
        else:
            set_run_font(run, size_pt=12, color=DARK_BLUE)
    return p

# ── Helper: add heading ─────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size_pt=16, bold=True, color=DARK_BLUE)
    else:
        set_run_font(run, size_pt=14, bold=True, color=DARK_BLUE)
    return p

# ── Build Document ──────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # Page margins
    sections = doc.sections[0]
    sections.top_margin = Cm(2.54)
    sections.bottom_margin = Cm(2.54)
    sections.left_margin = Cm(3.17)
    sections.right_margin = Cm(3.17)

    # ── Title ───────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run("年薪30万后才明白的职场真相：能力不是最重要的，重要的是它")
    set_run_font(run, size_pt=25, bold=True, color=DARK_BLUE)

    # Top border on first paragraph (subtitle/tagline area)
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(4)
    subtitle_p.paragraph_format.space_after = Pt(4)
    # Bottom border red
    pPr = subtitle_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="E74C3C"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    run = subtitle_p.add_run("职场能见度：决定你能走多远的隐形杠杆")
    set_run_font(run, size_pt=21, bold=True, color=DARK_BLUE)

    # Tagline
    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_before = Pt(8)
    tag_p.paragraph_format.space_after = Pt(16)
    run = tag_p.add_run("不是你不够努力，是你没有被看见")
    set_run_font(run, size_pt=18, bold=True, color=RED)

    add_divider(doc)

    # ── Opening Quote ───────────────────────────────────────────────────
    add_quote_block(doc, "在职场里，埋头苦干的人往往在原地打转，而懂得让自己被看见的人，才能撬动更大的世界。这不是功利，而是生存的法则。")

    add_divider(doc)

    # ── Chapter 1 ───────────────────────────────────────────────────────
    add_heading(doc, '一、我的30万年薪之路：从"技术宅"到"被看见的人"')

    add_body_with_highlight(doc, [
        ('五年前，我还是一个典型的"', False),
        ('技术宅"', True),
        ('——每天最早到公司，最晚离开，代码写得漂亮，bug修得飞快。我以为，只要能力足够强，升职加薪就是水到渠成的事。然而现实给了我狠狠的一巴掌：连续两年绩效评优，涨薪幅度却不到5%；同期入职的同事，能力明明不如我，却已经升了两级。', False),
    ])

    add_body_with_highlight(doc, [
        ('那时的我，', False),
        ('极度困惑"', True),
        ('甚至一度怀疑公司是不是在打压老实人。直到有一次，部门总监在会议上公开表扬了一个项目方案，而这个方案的核心思路其实是我提出来的——只不过，我把它写在了邮件里发给了小组长，小组长稍作整理后在汇报中呈现，功劳便与他绑定了。', False),
    ])

    add_body_with_highlight(doc, [
        ('那一刻我突然意识到：', False),
        ('在职场，能力只是入场券，能见度才是加速器。"', True),
        ('你的方案再出色，如果只有你的直属领导知道，它的价值就被局限在了一个很小的圈子里。', False),
    ])

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img1.jpg"), "职场中的能见度，往往比能力本身更能决定你的高度")

    add_divider(doc)

    # ── Chapter 2 ───────────────────────────────────────────────────────
    add_heading(doc, "二、能力陷阱：为什么越能干的人越容易遇到天花板")

    add_body_with_highlight(doc, [
        ('我见过太多', False),
        ('能力很强却始终在基层徘徊"', True),
        ('的人。他们有一个共同的特点：只关注"把事情做好"，却从不思考"让谁看到我把事情做好了"。', False),
    ])

    add_quote_block(doc, "职场不是学校，没有老师会主动发现你的才华。你的价值，需要通过有效的沟通和展示，才能被决策层感知到。")

    add_body_with_highlight(doc, [
        ('有个前同事小林，技术功底扎实，一个人能扛三个人的活。但他在会议上从不发言，汇报时只念数据，从不讲背后的思考逻辑。领导对他的评价永远是"踏实可靠"，但"缺乏领导力"、"不适合带团队"。三年后，他依然是高级工程师，而那些能力稍逊但善于表达、善于在关键时刻露脸的人，纷纷走上了管理岗。', False),
    ])

    add_body_with_highlight(doc, [
        ('这不是个例。', False),
        ('哈佛商学院的一项研究"', True),
        ('表明，在职场中，', False),
        ('个人能见度对晋升速度的影响权重高达40%"', True),
        ('，远超纯技术能力的25%。换句话说，你有多强并不重要，重要的是"别人觉得你有多强"。', False),
    ])

    add_highlight_box(doc, "能力决定你的下限，能见度决定你的上限。当你觉得怀才不遇时，先问问自己：有多少人真正知道你的才华？")

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img2.jpg"), "团队协作中，善于沟通的人更容易获得资源和机会")

    add_divider(doc)

    # ── Chapter 3 ───────────────────────────────────────────────────────
    add_heading(doc, '三、什么是真正的"职场能见度"')

    add_body_with_highlight(doc, [
        ('很多人误解了"能见度"，以为就是拍马屁、搞关系。真正的职场能见度，是', False),
        ('让你的工作成果、思考深度和个人品牌，被对的人以正确的方式感知到"', True),
        ('。它包含三个维度：', False),
    ])

    add_body_with_highlight(doc, [
        ('', False),
    ])

    add_body_with_highlight(doc, [
        ('', False),
    ])

    # Use a table for structured content
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Set column widths
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(12)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "维度"
    hdr_cells[1].text = "说明"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size_pt=12, bold=True, color=DARK_BLUE)
        # Shade header
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F8F9FA"/>')
        tcPr.append(shd)

    rows_data = [
        ("成果可见", "不仅做了，还要让相关方知道你做成了什么，解决了什么难题，创造了什么价值。"),
        ("观点可见", "在关键会议上有自己的声音，能提出独到见解，而不是永远附和别人。"),
        ("人格可见", "让同事和领导记住你是一个怎样的人——可靠、有想法、有担当。"),
    ]

    for dim, desc in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = dim
        row_cells[1].text = desc
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size_pt=11, color=DARK_BLUE)
            cell.vertical_alignment = 1  # CENTER vertically

    doc.add_paragraph()  # spacing

    add_body_with_highlight(doc, [
        ('这三个维度层层递进。', False),
        ('成果可见是基础，观点可见是升华，人格可见是长期资产。"', True),
        ('当你在这三个维度上都建立了存在感，你的职场天花板自然会随之抬高。', False),
    ])

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img3.jpg"), "在会议中主动表达观点，是提升职场能见度的重要方式")

    add_divider(doc)

    # ── Chapter 4 ───────────────────────────────────────────────────────
    add_heading(doc, '四、那些被忽视的"能人"： Visibility 缺失的代价')

    add_body_with_highlight(doc, [
        ('职场里从不缺少有能力的人，缺少的是', False),
        ('有能力且被看见的人"', True),
        ('。分享两个让我印象深刻的真实故事。', False),
    ])

    add_body_with_highlight(doc, [
        ('故事一：', False),
        ('"隐形冠军"的失落"', True),
    ])

    add_body_with_highlight(doc, [
        ('老张是一家互联网公司的后端架构师，十年工龄，技术栈深厚，是公司里少数能搞定高并发难题的人。但每次项目复盘，他总是坐在角落，让产品经理和前端负责人汇报。他的贡献被稀释在"团队努力"四个字里。去年公司裁员，老张在名单上——不是因为能力不行，而是高层根本不知道他是谁，裁掉他没有"政治成本"。', False),
    ])

    add_body_with_highlight(doc, [
        ('故事二：', False),
        ('"汇报高手"的逆袭"', True),
    ])

    add_body_with_highlight(doc, [
        ('小李入职时只是个普通运营，能力中等偏上。但她有一个习惯：每周给直属领导和跨部门协作的同事发一封', False),
        ('"本周亮点"邮件"', True),
        ('，简洁明了地列出自己完成的关键事项、数据成果和下周计划。半年后，全公司都知道有个"做事靠谱、思路清晰"的运营。一年后，她被破格提拔为运营组长，薪资涨幅超过50%。', False),
    ])

    add_quote_block(doc, "同样的能力，不同的能见度，结局天差地别。职场从来不是纯能力竞技场，而是能力与能见度共同作用的复杂系统。")

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img4.jpg"), "定期总结和汇报，是让自己被看见的最有效方法之一")

    add_divider(doc)

    # ── Chapter 5 ───────────────────────────────────────────────────────
    add_heading(doc, "五、如何系统提升你的职场能见度：可落地的行动清单")

    add_body_with_highlight(doc, [
        ('提升能见度不是让你变成"职场戏精"，而是', False),
        ('用专业和真诚的方式，让你的价值被更多人感知"', True),
        ('。以下是我实践多年总结的行动清单：', False),
    ])

    add_highlight_box(doc, '1. 建立"工作日志"习惯：每天记录做了什么、解决了什么、学到了什么。每周精选3条亮点，通过邮件或周报同步给直属领导和关键协作方。')

    add_highlight_box(doc, "2. 在会议上主动发言：哪怕只有一句话，也要让别人知道你在思考。可以从小处开始，比如补充一个数据、提出一个风险点、分享一个行业案例。")

    add_highlight_box(doc, "3. 跨部门建立连接：不要只和自己组的人熟。主动参加其他部门的分享会，在内部论坛发表专业见解，让更多人知道你的存在和专业领域。")

    add_highlight_box(doc, '4. 学会"向上管理"：定期与领导一对一沟通，不只是汇报进度，更要同步你的职业规划、学习成果和对团队的思考。让领导看到你的成长意愿。')

    add_highlight_box(doc, '5. 打造个人标签：找到你最擅长的领域，持续深耕并输出。比如"数据分析高手"、"用户增长专家"、"项目管理达人"。当别人有相关需求时，第一个想到你。')

    add_body_with_highlight(doc, [
        ('这些方法听起来简单，但', False),
        ('坚持三个月，你就会感受到明显的变化"', True),
        ('。你会发现领导开始主动找你讨论重要项目，同事遇到难题会先来请教你的意见，跨部门协作时你的名字自带信任背书。', False),
    ])

    add_image_with_caption(doc, os.path.join(IMG_DIR, "img5.jpg"), "持续输出专业见解，打造属于你的职场个人品牌")

    add_divider(doc)

    # ── Chapter 6 ───────────────────────────────────────────────────────
    add_heading(doc, "六、写在最后：能力与能见度，缺一不可")

    add_body_with_highlight(doc, [
        ('回到文章开头的问题：能力不是最重要的，重要的是什么？我的答案是——', False),
        ('职场能见度"', True),
        ('。但请一定要明白，能见度不能脱离能力单独存在。', False),
    ])

    add_quote_block(doc, "没有能力支撑的能见度，是空中楼阁；没有能见度加持的能力，是深埋地下的金矿。唯有二者兼备，才能在职场中走得更远、更稳。")

    add_body_with_highlight(doc, [
        ('年薪30万不是终点，而是一个新的起点。当你跨过这个门槛，你会发现，', False),
        ('真正决定你能走多远的，从来不是你能做多少事，而是有多少人知道你做了多少事，以及他们有多信任你去做更大的事。"', True),
    ])

    add_body_with_highlight(doc, [
        ('愿每一个努力的人，都能被世界温柔以待；更愿每一个有能力的人，都能学会让自己被看见。', False),
    ])

    add_divider(doc)

    # ── End Tags ────────────────────────────────────────────────────────
    add_para(doc, "#职场成长 #个人品牌 #职业发展 #能见度管理 #升职加薪",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=GRAY, space_before=8, space_after=8)

    add_para(doc, "——  END  ——",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=DARK_BLUE, space_before=8, space_after=8)

    # ── Save ────────────────────────────────────────────────────────────
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")

    # Character count
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text
    print(f"Total character count (approx): {len(full_text)}")
    return len(full_text)


if __name__ == "__main__":
    build_document()
