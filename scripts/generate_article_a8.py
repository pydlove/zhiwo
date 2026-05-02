#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a docx article in 清新马卡龙风 style.
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
IMG_DIR = "/tmp/article_a8_imgs"
OUT_DIR = "/Users/panyong/aio_project/小程序/docs/文章/2026-05-01/不离"
OUT_DOCX = os.path.join(OUT_DIR, "微信回\"好的\"的人，正在悄悄失去领导的信任.docx")

os.makedirs(OUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
PINK = RGBColor(0xF4, 0xA6, 0xB3)
TEAL_GREEN = RGBColor(0x5B, 0x8C, 0x85)
LIGHT_TEAL = RGBColor(0x7A, 0xAF, 0xA5)
GRAY = RGBColor(0xA0, 0xAE, 0xC0)
DARK_GRAY = RGBColor(0x71, 0x80, 0x96)
PURPLE_BORDER = RGBColor(0xC9, 0xB1, 0xFF)
LIGHT_PURPLE_BG = "FAF8FF"
MINT_BG = "F0FAF7"
CREAM_BG = "FFF8F0"
LIGHT_GRAY_BORDER = "E2E8F0"
DARK_TEXT = RGBColor(0x4A, 0x55, 0x68)

# ---------------------------------------------------------------------------
# Font helpers
# ---------------------------------------------------------------------------
def set_run_font(run, font_name="微软雅黑", size_pt=None, bold=False, italic=False, color=None):
    font = run.font
    if size_pt:
        font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), font_name)

# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------
def add_empty_paragraph(doc, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def set_paragraph_spacing(p, before=0, after=0, line_spacing=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_english_label(doc, text="WORKPLACE TIPS"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, font_name="Georgia", size_pt=10, color=PINK)
    set_paragraph_spacing(p, before=6, after=4)
    return p

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=28, bold=True, color=TEAL_GREEN)
    set_paragraph_spacing(p, before=4, after=6)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=18, bold=True, color=LIGHT_TEAL)
    set_paragraph_spacing(p, before=4, after=6)
    return p

def add_tagline(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=13, italic=True, color=GRAY)
    set_paragraph_spacing(p, before=4, after=10)
    return p

def add_body_paragraph(doc, text, bold_phrases=None):
    """
    Add a body paragraph. If bold_phrases is a list of substrings,
    those substrings will be rendered in bold PINK.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not bold_phrases:
        run = p.add_run(text)
        set_run_font(run, size_pt=12, color=DARK_GRAY)
    else:
        remaining = text
        for phrase in bold_phrases:
            idx = remaining.find(phrase)
            if idx == -1:
                continue
            if idx > 0:
                run = p.add_run(remaining[:idx])
                set_run_font(run, size_pt=12, color=DARK_GRAY)
            run = p.add_run(phrase)
            set_run_font(run, size_pt=12, bold=True, color=PINK)
            remaining = remaining[idx + len(phrase):]
        if remaining:
            run = p.add_run(remaining)
            set_run_font(run, size_pt=12, color=DARK_GRAY)
    set_paragraph_spacing(p, before=4, after=6, line_spacing=1.5)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("·  ·  ·")
    set_run_font(run, size_pt=14, color=PINK)
    set_paragraph_spacing(p, before=12, after=12)
    return p

def add_image_with_caption(doc, img_path, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    set_paragraph_spacing(p, before=8, after=4)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"▲ {caption_text}")
    set_run_font(run, size_pt=10, italic=True, color=GRAY)
    set_paragraph_spacing(cap, before=2, after=8)
    return p, cap

def add_quote_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    set_run_font(run, size_pt=13, color=DARK_GRAY)
    set_paragraph_spacing(p, before=10, after=10, line_spacing=1.5)

    # Top border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="{LIGHT_GRAY_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{LIGHT_GRAY_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Background shading via table trick (paragraph shading)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{MINT_BG}" w:val="clear"/>')
    pPr.append(shd)
    return p

def add_highlight_box1(doc, text):
    """Purple border, light purple bg, ✿ pink prefix, bold dark text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(f"✿  {text}")
    set_run_font(run, size_pt=12, bold=True, color=DARK_TEXT)
    set_paragraph_spacing(p, before=10, after=10, line_spacing=1.5)

    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="{PURPLE_BORDER}"/>'
        f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="{PURPLE_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{PURPLE_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="1" w:color="{PURPLE_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_PURPLE_BG}" w:val="clear"/>')
    pPr.append(shd)
    return p

def add_highlight_box2(doc, text):
    """Pink border, cream bg, bold teal text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size_pt=12, bold=True, color=TEAL_GREEN)
    set_paragraph_spacing(p, before=10, after=10, line_spacing=1.5)

    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="F4A6B3"/>'
        f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="F4A6B3"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="F4A6B3"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="1" w:color="F4A6B3"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CREAM_BG}" w:val="clear"/>')
    pPr.append(shd)
    return p

def add_chapter_marker(doc, number, title):
    """Number in 24pt bold pink, / in 16pt light gray, title in 14pt bold teal."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run(str(number))
    set_run_font(run1, size_pt=24, bold=True, color=PINK)
    run2 = p.add_run(" / ")
    set_run_font(run2, size_pt=16, color=GRAY)
    run3 = p.add_run(title)
    set_run_font(run3, size_pt=14, bold=True, color=TEAL_GREEN)
    set_paragraph_spacing(p, before=16, after=4)

    # Bottom thin border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{LIGHT_GRAY_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_end_tags(doc, tags):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(" ".join([f"#{t}" for t in tags]))
    set_run_font(run, size_pt=11, color=GRAY)
    set_paragraph_spacing(p, before=16, after=8)
    return p

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Top border (pink bottom border on first empty paragraph)
top_p = add_empty_paragraph(doc, space_after=2)
top_pPr = top_p._element.get_or_add_pPr()
top_pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="F4A6B3"/>'
    f'</w:pBdr>'
)
top_pPr.append(top_pBdr)

add_english_label(doc, "WORKPLACE TIPS")
add_title(doc, "微信回\"好的\"的人，正在悄悄失去领导的信任")
add_subtitle(doc, "职场沟通中的隐形陷阱与破局之道")
add_tagline(doc, "一条消息背后的职业素养，往往决定了你在领导心中的分量")

add_image_with_caption(doc, os.path.join(IMG_DIR, "img1.jpg"), "职场沟通，从一条微信消息开始")

add_body_paragraph(doc,
    "在快节奏的职场环境中，微信已经成为上下级沟通的重要工具。然而，很多人并没有意识到，"
    "自己随手回复的一个\"好的\"，可能正在悄悄消耗领导对自己的信任。",
    bold_phrases=["\"好的\""])

add_body_paragraph(doc,
    "这种看似礼貌的回应，实际上传递的信息非常有限。领导发出一条工作安排，"
    "期待的不仅仅是确认收到，更希望看到你对任务的理解、执行的思路以及完成的时间节点。"
    "一个干瘪的\"好的\"，就像一张空白支票，让领导心里没底。",
    bold_phrases=["\"好的\""])

add_divider(doc)

add_chapter_marker(doc, "01", "为什么\"好的\"会让人失望")

add_body_paragraph(doc,
    "小张是一名工作三年的运营专员。某天下午，领导在微信上给他布置了一项紧急任务："
    "\"把上季度的数据整理一下，下班前发我。\"小张秒回了一个\"好的\"，然后继续忙手头的工作。",
    bold_phrases=["\"好的\""])

add_body_paragraph(doc,
    "两个小时后，领导再次发来消息：\"数据好了吗？\"小张这才恍然大悟，"
    "原来领导需要的是一份完整的分析报告，而不是简单的数字汇总。"
    "他手忙脚乱地赶工，之后在截止时间前勉强交差，但质量可想而知。")

add_quote_block(doc,
    "\"好的\"这两个字，表面上是对指令的服从，实际上却暴露了一个致命问题："
    "回复者没有进行有效思考，也没有展现出对任务的主动理解。")

add_body_paragraph(doc,
    "从心理学角度来看，领导在布置任务时，内心往往伴随着一定的焦虑感。"
    "他们需要确认下属听懂了、有能力做、并且能在预期时间内完成。"
    "一个高质量的回复，能够很大程度上缓解这种焦虑。而一个\"好的\"，"
    "则让这种焦虑悬在半空，甚至不断放大。",
    bold_phrases=["\"好的\""])

add_image_with_caption(doc, os.path.join(IMG_DIR, "img2.jpg"), "高质量的回复，是对领导焦虑的很好安抚")

add_highlight_box1(doc,
    "真正成熟的职场人，懂得在回复中传递确定性："
    "任务理解无误、执行有思路、节点可预期。")

add_divider(doc)

add_chapter_marker(doc, "02", "领导真正想看到什么样的回复")

add_body_paragraph(doc,
    "那么，面对领导的微信安排，怎样的回复才算合格呢？我们可以从三个维度来拆解："
    "确认信息、展示思路、明确节点。")

add_body_paragraph(doc,
    "首先，确认信息。在回复中简要复述任务要点，让领导知道你理解正确。"
    "例如：\"收到，您需要我上季度各渠道的数据汇总及趋势分析，对吗？\""
    "这种回复既体现了你的专注，也给领导一个纠偏的机会。",
    bold_phrases=["确认信息"])

add_body_paragraph(doc,
    "其次，展示思路。简单说明你打算怎么做，让领导看到你的专业度和条理性。"
    "例如：\"我打算先从后台导出原始数据，然后按渠道分类统计，之后做同比和环比分析。\""
    "这样的回复，会让领导觉得你做事有章法，值得信赖。",
    bold_phrases=["展示思路"])

add_body_paragraph(doc,
    "再者，明确节点。给出一个具体的时间承诺，让领导心里有数。"
    "例如：\"预计今天下午四点前完成初稿，发您邮箱。如有调整，我会提前同步。\""
    "明确的时间节点，是对双方效率的尊重。",
    bold_phrases=["明确节点"])

add_highlight_box2(doc,
    "一个优质的回复模板：收到+复述+思路+时间节点。"
    "例如：\"收到，我来整理上季度数据，先按渠道分类统计再做趋势对比，预计四点前发您。\"")

add_image_with_caption(doc, os.path.join(IMG_DIR, "img3.jpg"), "清晰的沟通，是职场协作的润滑剂")

add_divider(doc)

add_chapter_marker(doc, "03", "不同场景下的高情商回复示范")

add_body_paragraph(doc,
    "职场沟通并非千篇一律，不同的场景需要不同的回复策略。"
    "以下是几种常见场景下的高情商回复方式，供你参考。")

add_body_paragraph(doc,
    "场景一：接到常规任务。领导说：\"把这个方案优化一下。\""
    "低情商回复：\"好的。\"高情商回复：\"收到，我重点优化用户转化路径和成本测算部分，"
    "预计明天上午十点前给您一版，您看可以吗？\"",
    bold_phrases=["场景一"])

add_body_paragraph(doc,
    "场景二：接到模糊指令。领导说：\"近来竞品动作不少，你留意一下。\""
    "低情商回复：\"好的。\"高情商回复：\"明白，我本周内梳理三家主要竞品的近期动态，"
    "包括产品更新、营销活动和市场反馈，周五前给您一份简报。\"",
    bold_phrases=["场景二"])

add_body_paragraph(doc,
    "场景三：时间冲突无法承接。领导说：\"下午来开个会。\""
    "低情商回复：\"好的。\"（实际上已有安排）高情商回复：\"收到，我下午两点到四点有客户对接，"
    "如果会议可以调整到四点以后，我一定参加。或者我先看会议纪要，有需要在会后单独沟通？\"",
    bold_phrases=["场景三"])

add_body_paragraph(doc,
    "场景四：需要资源支持。领导说：\"这个活动你来牵头。\""
    "低情商回复：\"好的。\"高情商回复：\"感谢信任！为了确保活动效果，"
    "我需要设计同事配合出两套视觉方案，以及预算审批走加急流程。"
    "我先列一个需求清单，今天下午找您确认，可以吗？\"",
    bold_phrases=["场景四"])

add_quote_block(doc,
    "高情商的回复，核心在于换位思考：如果你是领导，"
    "看到这条消息，是否能放下心来？是否还有疑虑需要追问？")

add_image_with_caption(doc, os.path.join(IMG_DIR, "img4.jpg"), "换位思考，是高情商沟通的起点")

add_divider(doc)

add_chapter_marker(doc, "04", "长期积累，建立你的职场口碑")

add_body_paragraph(doc,
    "微信沟通虽然只是职场交往的一个切片，但它反映的是一个人的职业素养和工作习惯。"
    "每一次高质量的回复，都是在为自己的职场口碑添砖加瓦。")

add_body_paragraph(doc,
    "久而久之，领导会在潜意识里形成对你的稳定预期："
    "\"这件事交给小李，他一定能按时交付，而且质量有保障。\""
    "这种信任，是晋升、加薪、被委以重任的重要基础。")

add_body_paragraph(doc,
    "相反，如果你总是用\"好的\"\"收到\"\"明白\"来敷衍，"
    "领导可能会逐渐把你归为\"需要反复确认\"\"不太靠谱\"的那一类人。"
    "一旦这种印象固化，想要扭转就需要付出成倍的努力。",
    bold_phrases=["\"好的\"\"收到\"\"明白\""])

add_highlight_box1(doc,
    "职场没有白走的路，每一条认真回复的消息，"
    "都在悄悄为你打开更大的成长空间。")

add_body_paragraph(doc,
    "当然，高情商沟通不是让你写得像作文一样冗长。"
    "核心原则是：在简洁的前提下，尽可能多地传递有效信息。"
    "领导的时间也很宝贵，三句话能说清楚的事，不要写成三段。")

add_body_paragraph(doc,
    "另外，及时性同样重要。看到消息后尽快回复，"
    "哪怕只是一个简短的确认，也比已读不回要好得多。"
    "如果确实需要较长时间才能给出完整回复，可以先发一条："
    "\"收到，我先梳理一下，半小时内给您详细回复。\""
    "这样既表达了重视，也争取了思考时间。")

add_image_with_caption(doc, os.path.join(IMG_DIR, "img5.jpg"), "每一条认真的消息，都是职场进阶的阶梯")

add_highlight_box2(doc,
    "从今天开始，试着把\"好的\"换成更具体的表达。"
    "你会发现，领导对你的态度，正在发生微妙而积极的变化。")

add_body_paragraph(doc,
    "职场沟通是一门需要长期修炼的艺术。"
    "从一条微信消息开始，培养结构化表达的习惯，"
    "不仅能提升你在领导心中的专业形象，也能让你的工作更加高效有序。")

add_body_paragraph(doc,
    "除了回复内容本身，还有一些细节值得留意。"
    "比如，适当使用分段和标点，让消息更易读；"
    "避免在深夜或凌晨回复工作消息，以免给领导留下\"没有边界感\"的印象；"
    "如果领导发了语音，尽量用文字回复，方便对方查阅和转发。")

add_body_paragraph(doc,
    "此外，定期复盘自己的沟通方式也很重要。"
    "你可以每周花几分钟，回顾一下自己与领导的几次关键对话，"
    "思考哪些地方可以改进。这种自我迭代的意识，"
    "会让你在职场中越走越稳。")

add_body_paragraph(doc,
    "记住，真正优秀的职场人，从不把沟通当作负担，"
    "而是把它视为展示自己、连接他人的重要机会。"
    "下一次收到领导的消息时，不妨多花十秒钟，"
    "给自己一个更出色的回复。")

add_divider(doc)

add_end_tags(doc, ["职场沟通", "高情商回复", "微信礼仪", "职业发展", "向上管理"])

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT_DOCX)
print(f"Saved: {OUT_DOCX}")

# ---------------------------------------------------------------------------
# Copy images to output dir
# ---------------------------------------------------------------------------
out_img_dir = os.path.join(OUT_DIR, "images")
os.makedirs(out_img_dir, exist_ok=True)
for f in os.listdir(IMG_DIR):
    src = os.path.join(IMG_DIR, f)
    dst = os.path.join(out_img_dir, f)
    shutil.copy2(src, dst)
print(f"Images copied to: {out_img_dir}")

# ---------------------------------------------------------------------------
# Verify character count
# ---------------------------------------------------------------------------
from docx import Document as DocReader
doc2 = DocReader(OUT_DOCX)
full_text = []
for para in doc2.paragraphs:
    full_text.append(para.text)
total_chars = sum(len(t) for t in full_text)
print(f"Total characters: {total_chars}")

# ---------------------------------------------------------------------------
# Forbidden words check
# ---------------------------------------------------------------------------
forbidden_map = {
    "永久": "长期", "关注": "留意", "最新": "新近", "最": "（避免使用，改用\"很\"\"非常\"等）",
    "最大": "很大", "群": "圈子/团体", "第一": "领先", "唯一": "少有的", "顶级": "优质",
    "最高级": "更高级", "万能": "多用途", "100%": "绝大多数", "绝对": "相当", "最好": "很好",
    "最佳": "很出色", "最强": "很强", "首选": "优选", "全网最低": "价格优惠", "最便宜": "性价比高",
    "国家级": "省级以上", "首家": "率先", "独家": "特有", "史上最": "非常", "特效": "效果明显",
    "震惊": "令人关注", "跳楼价": "超值价", "亏本卖": "薄利多销", "错过再等一年": "限时优惠",
    "保证治愈": "有望改善", "药到病除": "有助缓解", "治愈": "改善", "根治": "从根本上改善",
    "无副作用": "成分温和", "疗效": "效果", "秘方": "传统方法", "偏方": "传统方法",
    "神医": "资深医师", "神药": "有效药物", "无效退款": "可咨询售后", "抗癌": "辅助调理",
    "增高": "促进发育", "一吃就瘦": "配合运动", "减肥神药": "辅助产品", "延寿": "健康养生",
    "几天见效": "坚持服用", "暴富": "财富增长", "保本保息": "稳健型", "稳赚不赔": "收益相对稳定",
    "保本": "稳健型", "稳赚": "收益相对稳定", "零风险": "风险可控", "原始股": "股权投资",
    "内幕消息": "市场分析", "配资": "融资服务", "收益率100%": "预期收益", "高回报": "预期收益较好",
    "假一赔十": "正品保障", "免费送": "赠品活动", "烟草": "其他商品", "爆炸性": "重要",
    "突发": "刚刚", "重磅": "重要", "速看": "值得关注", "躺赚": "被动收入", "限时免费": "免费体验",
    "零元购": "免费试用", "全网首发": "率先发布", "白菜价": "性价比高", "先到先得": "数量有限",
    "转发此文章": "欢迎分享", "点击关注": "欢迎关注", "最后机会": "机会难得", "算命": "运势分析",
    "改运": "积极心态", "约炮": "交友", "一夜情": "短期关系", "赌博": "娱乐", "彩票预测": "彩票分析",
    "时时彩": "彩票", "外挂": "辅助工具", "走私": "跨境贸易", "破解版": "正版授权",
    "盗版": "非正版", "VPN": "网络工具", "翻墙": "跨境访问", "枪支": "器械", "代购": "海外购",
    "占卜": "性格测试", "八字": "出生日期", "风水": "环境布局", "开光": "仪式", "写真": "照片",
    "命格": "性格特点", "转运": "好运气", "荐股": "投资建议", "电子烟": "雾化器", "百分百": "绝大多数"
}

forbidden_phrases = ["分享到朋友圈", "不转不是中国人", "不转发死全家", "转发好运", "福利姬", "无码", "有码"]

all_text = "\n".join(full_text)
found = []
for bad, good in forbidden_map.items():
    if bad in all_text:
        found.append((bad, good))
for bad in forbidden_phrases:
    if bad in all_text:
        found.append((bad, "REMOVE"))

if found:
    print("WARNING: Found forbidden words:")
    for bad, good in found:
        print(f"  {bad} -> {good}")
else:
    print("No forbidden words found. PASS.")
