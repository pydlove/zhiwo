from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

output_dir = "/Users/panyong/aio_project/小程序/创作技巧"
os.makedirs(output_dir, exist_ok=True)

# 模板数据 + 配色方案
templates = [
    {
        "name": "精美商务版",
        "theme": "商务蓝",
        "primary": (30, 58, 95),      # 深蓝 #1e3a5f
        "accent": (0, 112, 192),      # 亮蓝 #0070c0
        "intro": "适用于商业分析、行业洞察、品牌策划、企业管理等专业内容。结构清晰、逻辑严密，强调观点输出与 actionable insights。",
        "structure": [
            ("一、核心观点", "在开篇用简洁有力的语言点明文章主旨，让读者3秒内get到价值。"),
            ("二、背景分析", "结合行业趋势、市场数据或宏观环境，交代问题的来龙去脉。"),
            ("三、深度解读", "多角度剖析，引用案例和数据支撑论点，体现专业度。"),
            ("四、行动建议", "给出具体、可落地的解决方案，帮助读者把认知转化为行动。"),
            ("五、结语", "总结升华，用一句有力的话收尾，强化记忆点。")
        ],
        "example_title": "示例：2024年内容营销的三大趋势与应对策略",
        "example_sections": [
            ("一、核心观点", "2024年，内容营销正在从'流量驱动'转向'价值驱动'。品牌需要更加注重深度内容、用户关系与长期信任的建立。"),
            ("二、背景分析", "随着平台算法日益成熟，用户对低质内容的免疫力增强，单纯依靠标题党和投放获取流量的模式已难以为继。同时，AI生成内容的爆发也让'人味儿'和'专业度'成为稀缺资源。"),
            ("三、深度解读", "趋势一：短视频+长图文融合。用户既需要快速获取信息，也渴望深度学习。趋势二：私域运营成为标配。公域获客成本持续走高，私域的复购和转介绍价值凸显。趋势三：AI辅助创作普及。善用AI的团队效率提升3-5倍，但核心创意仍需人类把关。"),
            ("四、行动建议", "1. 建立内容中台，统一选题、生产和分发流程；2. 培养兼具行业洞察与AI工具使用能力的内容团队；3. 制定私域运营SOP，将公域流量沉淀为可反复触达的用户资产。"),
            ("五、结语", "内容营销的终局不是比拼谁的声音更大，而是比拼谁能为用户创造真正的价值。先人一步转型，就能在竞争中赢得先机。")
        ],
        "tips": [
            "核心观点段控制在100字以内，避免绕弯子。",
            "背景分析尽量引用权威数据或最新行业报告。",
            "行动建议使用清单体，方便读者截图收藏。",
            "结语要有一句'金句'，增强传播性。"
        ]
    },
    {
        "name": "简约资讯版",
        "theme": "资讯灰",
        "primary": (55, 65, 81),      # 深灰 #374151
        "accent": (107, 114, 128),    # 中灰 #6b7280
        "intro": "适用于热点追踪、新闻速览、行业周报、事件复盘等内容。风格客观中立，信息密度高，适合快速阅读。",
        "structure": [
            ("导语", "用一两句话概括事件核心，快速抓住注意力。"),
            ("事件回顾", "按时间线或逻辑顺序梳理经过，突出关键节点。"),
            ("关键数据", "用数字说话，列出与话题相关的核心统计。"),
            ("多方观点", "汇总业内人士、专家或网友看法，呈现多元视角。"),
            ("结语", "简要总结影响或未来走向，保持客观。")
        ],
        "example_title": "示例：OpenAI发布GPT-4o：多模态能力与免费策略引发行业震动",
        "example_sections": [
            ("导语", "5月14日，OpenAI正式发布新一代旗舰模型GPT-4o，支持文本、音频、图像任意组合输入输出，并将向免费用户开放部分能力。"),
            ("事件回顾", "2022年11月，ChatGPT横空出世，掀起全球AI浪潮。2023年3月，GPT-4发布，多模态能力初现。2024年5月，GPT-4o将响应速度提升至接近人类对话水平，并推出桌面端应用。"),
            ("关键数据", "GPT-4o语音模式平均响应时间320毫秒；支持50种语言，覆盖全球97%的人口；免费用户的消息限额是GPT-4的5倍。"),
            ("多方观点", "业内专家认为，GPT-4o的免费策略将加速AI普及，同时对中小模型厂商形成巨大压力。部分用户担忧，更强的多模态能力可能带来深度伪造等伦理风险。"),
            ("结语", "GPT-4o的发布标志着人机交互进入新阶段。对于内容创作者而言，拥抱AI工具、提升不可替代的创意能力，是应对变革的最佳策略。")
        ],
        "tips": [
            "导语遵循'5W1H'原则，快速交代关键信息。",
            "事件回顾避免冗长，只保留对理解事件至关重要的节点。",
            "数据要标注来源，增强可信度。",
            "多方观点注意平衡，避免单一立场。"
        ]
    },
    {
        "name": "情感故事版",
        "theme": "暖橙",
        "primary": (194, 65, 12),     # 暖橙 #c2410c
        "accent": (234, 88, 12),      # 亮橙 #ea580c
        "intro": "适用于情感共鸣类、人生感悟类、成长励志类内容。以故事为主线，通过具体场景和细腻描写打动读者。",
        "structure": [
            ("故事/事件一", "用一个具体的生活片段引出第一个观点或感悟。"),
            ("故事/事件二", "递进展开，通过第二个片段深化主题。"),
            ("故事/事件三", "进一步拓展，形成层层递进的情感张力。"),
            ("总结升华", "从个人经历上升到普遍认知，引发读者共鸣。")
        ],
        "example_title": "示例：人到中年，才明白的三件事",
        "example_sections": [
            ("第一件事：健康是一切的基础", "没有健康，所有的成就、财富、地位都不过是空中楼阁。中年以后，身体机能开始走下坡路，曾经熬夜通宵第二天依然生龙活虎的日子一去不复返。这个时候，规律作息、适度运动、定期体检，不是选择题，而是必答题。"),
            ("第二件事：关系需要经营", "无论是亲情、友情还是爱情，再好的关系也经不起长期的忽视和消耗。中年人的社交圈往往会自动收缩，留下的都是真正值得珍惜的人。用心经营这些关系，才能在关键时刻有所依靠。"),
            ("第三件事：心态决定状态", "中年危机并不可怕，可怕的是你对此毫无准备，或者被焦虑彻底吞噬。学会接纳自己的不完美，接纳生活的不如意，保持一颗平和而积极的心，才能在这个阶段活出新的精彩。"),
            ("", "人生下半场，拼的不是谁跑得快，而是谁走得稳。愿你早日明白这三件事，把接下来的路走得更加从容。")
        ],
        "tips": [
            "故事要具体，少用抽象道理，多用场景和细节。",
            "每个故事片段控制在200-300字，保持阅读节奏。",
            "善用第一人称或贴近读者的视角，拉近距离。",
            "结尾升华要克制，点到为止，给读者留白。"
        ]
    },
    {
        "name": "科技专业版",
        "theme": "科技青",
        "primary": (37, 99, 235),     # 科技蓝 #2563eb
        "accent": (14, 165, 233),     # 青色 #0ea5e9
        "intro": "适用于数码测评、产品评测、技术解读、软硬件推荐等内容。风格理性客观，注重实测数据和购买决策指导。",
        "structure": [
            ("产品概述", "介绍品牌、型号、定位、主要卖点，建立基础认知。"),
            ("外观设计", "描述外观、材质、尺寸重量，结合主观感受评价。"),
            ("性能实测", "通过测试数据或对比，客观呈现关键性能指标。"),
            ("优缺点分析", "理性分析优势与不足，帮助读者全面了解。"),
            ("购买建议", "针对不同需求给出是否值得购买及配置建议。")
        ],
        "example_title": "示例：iPhone 15 Pro 深度评测： titanium机身+A17 Pro，值不值得买？",
        "example_sections": [
            ("产品概述", "iPhone 15 Pro是苹果2023年秋季发布的旗舰机型，首次采用钛金属边框，搭载A17 Pro芯片，USB-C接口取代沿用多年的Lightning。定位高端商务和极客用户。"),
            ("外观设计", "钛金属中框带来了显著的减重效果，整机重量从206g降至187g。边缘经过微弧处理，握持感比前代更舒适。提供原色钛金属、蓝色钛金属等四款配色，质感高级但容易沾染指纹。"),
            ("性能实测", "A17 Pro在Geekbench 6单核跑分约2900，多核约7200，GPU性能提升约20%。《原神》最高画质30分钟平均帧率59.2fps，机身温度控制在45°C以内。USB 3.0传输速度可达10Gbps，专业视频 workflow 效率大幅提升。"),
            ("优缺点分析", "优点：极致性能、钛金属减重、USB-C通用性强、Action Button可自定义。缺点：长焦仍为3倍光学变焦，望远能力不及竞品；高负载下发热依然明显；价格门槛较高。"),
            ("购买建议", "如果你是iPhone 13及更早机型的用户，15 Pro的升级感知非常明显，值得购买。如果你已经是14 Pro用户，除非对USB-C和钛金属有强烈需求，否则可以再等一代。存储建议至少选择256GB版本。")
        ],
        "tips": [
            "产品概述段控制在150字以内，快速建立认知。",
            "实测数据要注明测试环境和工具，增强可信度。",
            "优缺点分析要敢于指出不足，避免像软文。",
            "购买建议按人群细分（学生/商务/极客/普通用户）。"
        ]
    },
    {
        "name": "娱乐轻松版",
        "theme": "活力粉",
        "primary": (219, 39, 119),     # 玫红 #db2777
        "accent": (236, 72, 153),      # 亮粉 #ec4899
        "intro": "适用于娱乐八卦、热点玩梗、轻松盘点、幽默故事等内容。文风活泼俏皮，追求高互动和强传播。",
        "structure": [
            ("开篇引入", "用热点梗、轻松话题或趣味开场白拉近距离。"),
            ("正文展开", "围绕主题加入趣闻轶事、神评论或搞笑桥段。"),
            ("亮点盘点", "提炼吸睛看点或槽点，逐一解读激发共鸣。"),
            ("结尾升华", "用俏皮话、反转或温情收尾，引导互动。")
        ],
        "example_title": "示例：这届网友是懂评论的，十大神回复笑不活了",
        "example_sections": [
            ("开篇引入", "都说自古评论出人才，今天带大家一起见识一下，什么叫'评论区比正文好看'。温馨提示：吃饭喝水时请勿阅读，否则后果自负。"),
            ("正文展开", "有人问：为什么程序员总是穿格子衫？神回复：因为代码写错了，可以直接从衣服上找到对应的格子定位bug。还有人问：如何委婉地拒绝别人借钱？热评第一：我最近也在等人还钱，要不你排个号？"),
            ("亮点盘点", "盘点本期最佳神评论：最扎心奖——'我的工资增长速度，连Wi-Fi信号都不如'；最清醒奖——'不要因为5分钟的不开心，就浪费23小时55分钟的开心'；最真实奖——'小时候以为早睡早起身体好是一句口号，长大后才发现这是三个愿望'。"),
            ("结尾升华", "生活已经够累了，偶尔在评论区笑笑，也是一种低成本的幸福。如果你也被逗乐了，欢迎在评论区留下你见过的神回复，让我们一起哈哈哈哈！")
        ],
        "tips": [
            "开篇3句话内必须出现笑点或梗，抓住注意力。",
            "正文多使用短句、对话体和网友口吻。",
            "亮点盘点用'奖项''排名''榜单'形式，增强趣味性。",
            "结尾一定要引导评论/点赞/转发，提升互动率。"
        ]
    }
]


def create_docx(tpl, filepath):
    doc = Document()
    primary = RGBColor(*tpl['primary'])
    accent = RGBColor(*tpl['accent'])

    # 标题
    title = doc.add_heading(f"样式模板：{tpl['name']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = primary

    # 副标题 / 主题标签
    tag = doc.add_paragraph(f"主题风格：{tpl['theme']}")
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tag.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = accent
        run.font.italic = True
    tag.paragraph_format.space_after = Pt(12)

    # 导语区块（带颜色强调）
    intro_label = doc.add_paragraph()
    intro_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label_run = intro_label.add_run("【导语】")
    label_run.bold = True
    label_run.font.size = Pt(12)
    label_run.font.color.rgb = primary
    intro_label.paragraph_format.space_after = Pt(4)

    intro_text = doc.add_paragraph(tpl['intro'])
    intro_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_text.paragraph_format.line_spacing = 1.5
    intro_text.paragraph_format.left_indent = Inches(0.2)
    for run in intro_text.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(55, 65, 81)

    # 结构解析
    heading = doc.add_heading("一、模板结构解析", level=2)
    for run in heading.runs:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = primary

    for section_title, section_desc in tpl['structure']:
        h = doc.add_heading(section_title, level=3)
        for run in h.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = accent
        p = doc.add_paragraph(section_desc)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(55, 65, 81)

    # 示例文章
    doc.add_paragraph()
    ex_heading = doc.add_heading("二、示例文章", level=2)
    for run in ex_heading.runs:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = primary

    ex_title = doc.add_paragraph(tpl['example_title'])
    ex_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ex_title.runs:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = primary
    ex_title.paragraph_format.space_before = Pt(6)
    ex_title.paragraph_format.space_after = Pt(10)

    for section_title, section_content in tpl['example_sections']:
        if section_title:
            h = doc.add_heading(section_title, level=3)
            for run in h.runs:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = accent
        p = doc.add_paragraph(section_content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(11)

    # 使用技巧
    doc.add_paragraph()
    tips_heading = doc.add_heading("三、使用技巧", level=2)
    for run in tips_heading.runs:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = primary

    for i, tip in enumerate(tpl['tips'], 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        num_run = p.add_run(f"{i}. ")
        num_run.font.size = Pt(11)
        num_run.font.bold = True
        num_run.font.color.rgb = accent
        text_run = p.add_run(tip)
        text_run.font.size = Pt(11)
        text_run.font.color.rgb = RGBColor(55, 65, 81)

    # 结语
    doc.add_paragraph()
    con_heading = doc.add_heading("结语", level=2)
    for run in con_heading.runs:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = primary

    con_p = doc.add_paragraph()
    con_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    con_p.paragraph_format.line_spacing = 1.5
    con_p.paragraph_format.left_indent = Inches(0.2)
    quote_run = con_p.add_run("「")
    quote_run.font.size = Pt(13)
    quote_run.font.bold = True
    quote_run.font.color.rgb = accent
    body_run = con_p.add_run(
        f"{tpl['name']}」是经过验证的高效内容结构。掌握它的用法，结合 AI 辅助创作，能够大幅提升写作效率和内容质量。建议您根据自身赛道特点，选择最适合的模板并设置为默认模板。"
    )
    body_run.font.size = Pt(11)
    body_run.font.italic = True
    body_run.font.color.rgb = RGBColor(55, 65, 81)

    doc.save(filepath)
    print(f"已生成: {filepath}")


for tpl in templates:
    filepath = os.path.join(output_dir, f"样式模板_{tpl['name']}.docx")
    create_docx(tpl, filepath)

print(f"\n全部完成，共生成 {len(templates)} 个 docx 文件，保存在: {output_dir}")
