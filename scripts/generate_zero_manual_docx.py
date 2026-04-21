from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_dir = "/Users/panyong/aio_project/小程序/创作技巧"
os.makedirs(output_dir, exist_ok=True)

doc = Document()

# 标题
title = doc.add_heading("零人工创作操作手册", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

# 导语
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro_run = intro.add_run("【导语】")
intro_run.bold = True
intro_run.font.size = Pt(12)
intro_run.font.color.rgb = RGBColor(0, 112, 192)
intro_text = doc.add_paragraph(
    "本手册面向平台创作者，详细介绍如何使用系统的「零人工创作」功能，从选题到成文，全程由 AI 自动完成，帮助您高效产出优质内容。"
)
intro_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro_text.runs[0].font.size = Pt(11)

sections = [
    ("一、功能概述", [
        "「零人工创作」是本平台的智能写作核心功能之一。与「半人工创作」不同，零人工模式只需提供文章标题，AI 即可根据标题、赛道特征以及平台调性，自动生成完整文章内容。",
        "适用场景：需要快速出稿、对选题已有明确方向、希望通过 AI 完成初稿再稍作润色的创作者。"
    ]),
    ("二、使用前的准备", [
        "1. 账号权限：由管理员在后台为您创建账号并分配平台权限（公众号 / 今日头条 / 百家号）和 AI 使用配额。",
        "2. 订阅赛道：在创作中心，您需要先订阅感兴趣的赛道（如情感故事、科技数码、职场成长等）。管理员可设置每位用户的订阅上限。",
        "3. 选择默认模板：在个人中心 → 样式模板中，可设置默认的文章结构模板（精美商务版、简约资讯版、情感故事版、科技专业版、娱乐轻松版）。零人工创作生成的文章将自动套用该模板框架。",
        "4. 账号有效期：请确保账号在有效期内。过期后将无法保存草稿、调用 AI 生成或导出 Word。"
    ]),
    ("三、零人工创作操作流程", [
        "步骤 1：进入创作中心",
        "在首页浏览各赛道数据，点击感兴趣的赛道卡片，或点击赛道详情页中的「✨ AI 生成创作方向」按钮，进入创作流程。",
        "",
        "步骤 2：选择赛道",
        "在创作页第一步，系统会展示您已订阅的赛道列表。点击目标赛道右侧的「进入创作」按钮。",
        "",
        "步骤 3：选择创作方向",
        "系统会展示该赛道在当前平台下的「每日推荐」主题（由运营人员在后台配置）。推荐主题包含标题和简要描述，点击任意推荐卡片即可自动带入标题。",
        "如果您已有明确标题，也可点击「✏️ 自定义创作」，手动输入标题。",
        "",
        "步骤 4：切换为零人工模式",
        "进入编辑器后，页面顶部左侧有两个模式按钮：「半人工创作」和「零人工创作」。点击「零人工创作」切换到全自动写作模式。",
        "",
        "步骤 5：AI 自动生成文章",
        "确认标题无误后，点击右上角的「✨ AI 生成大纲」按钮（在零人工模式下，AI 将基于标题和赛道特征生成完整文章内容，并自动填充到编辑器中）。",
        "生成过程中请保持网络畅通。生成完成后，您可直接在编辑器中查看全文。",
        "",
        "步骤 6：编辑与预览",
        "AI 生成的内容为初稿，您可以使用编辑器提供的富文本工具进行二次调整：加粗、斜体、H1/H2 标题、无序/有序列表、插入链接等。",
        "点击「实时预览」标签，可查看文章在移动端的真实渲染效果。"
    ]),
    ("四、模板说明", [
        "系统内置 5 种文章结构模板，每种模板对应不同的段落框架：",
        "• 精美商务版：核心观点 → 背景分析 → 深度解读 → 行动建议 → 结语",
        "• 简约资讯版：导语 → 事件回顾 → 关键数据 → 多方观点 → 结语",
        "• 情感故事版：多段式故事叙述 + 情感升华结尾",
        "• 科技专业版：产品概述 → 外观设计 → 性能实测 → 优缺点分析 → 购买建议",
        "• 娱乐轻松版：开篇引入 → 正文展开 → 亮点盘点 → 结尾升华",
        "",
        "建议在个人中心提前设置与您的赛道最匹配的模板，这样 AI 生成时会更好地贴合目标文风。"
    ]),
    ("五、保存与导出", [
        "保存草稿：点击编辑器右上角「保存草稿」按钮，系统会将当前标题、内容和创作模式一并保存。您可以在「我的草稿」或「我的创作」中继续编辑。",
        "导出 Word：文章定稿后，点击「导出 Word」按钮，即可将文章内容下载为 Word 文档，方便您进一步排版或提交发布。",
        "",
        "注意：每次 AI 生成都会计入您的 AI 使用配额（在个人中心可查看剩余额度）。配额耗尽后请联系管理员充值。"
    ]),
    ("六、常见问题", [
        "Q1：零人工创作和半人工创作有什么区别？",
        "A：半人工创作模式下，AI 仅生成文章大纲，正文需要您自行撰写；零人工创作模式下，AI 直接生成完整文章正文。",
        "",
        "Q2：AI 生成的内容可以直接发布吗？",
        "A：AI 生成的是高质量初稿，建议您在发布前进行事实核查、品牌调性调整和合规审查。",
        "",
        "Q3：为什么点击 AI 生成后没有反应？",
        "A：请检查以下情况：账号是否已过期、AI 配额是否已用完、网络连接是否正常、标题是否为空。",
        "",
        "Q4：可以修改已经生成的文章吗？",
        "A：可以。生成的内容会完整填充到编辑器中，您可以像编辑普通文档一样任意修改。",
        "",
        "Q5：每日推荐的数据从哪里来？",
        "A：每日推荐由平台运营人员在管理后台配置，按「平台 + 赛道」维度维护，确保推荐主题紧跟热点。"
    ]),
    ("七、创作建议", [
        "1. 标题越清晰，AI 生成效果越好。建议在自定义创作时，使用包含数字、痛点或场景的具体标题。",
        "2. 善用参考文章。编辑器左侧会展示当前赛道下的爆款参考文章链接，阅读这些文章有助于把握该赛道的读者偏好。",
        "3. 定期查看「帮助中心」和「创作技巧」板块，学习平台提供的运营干货，提升内容打开率和互动率。",
        "4. 保持创作节奏，建议每周固定时间登录平台选题、生成、润色，形成稳定的内容产出习惯。"
    ])
]

for section_title, paras in sections:
    heading = doc.add_heading(section_title, level=2)
    for run in heading.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    for p_text in paras:
        if p_text.strip() == "":
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(p_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(11)

# 结语
doc.add_paragraph()
conclusion_heading = doc.add_heading("结语", level=2)
for run in conclusion_heading.runs:
    run.font.size = Pt(14)
    run.font.bold = True

conclusion_p = doc.add_paragraph(
    "零人工创作是提升内容生产效率的利器。掌握本手册中的操作步骤和技巧，您就能快速将选题转化为优质文章。祝您创作顺利，篇篇爆款！"
)
conclusion_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
conclusion_p.paragraph_format.line_spacing = 1.5
for run in conclusion_p.runs:
    run.font.size = Pt(11)
    run.font.italic = True

filepath = os.path.join(output_dir, "零人工创作操作手册.docx")
doc.save(filepath)
print(f"已生成: {filepath}")
