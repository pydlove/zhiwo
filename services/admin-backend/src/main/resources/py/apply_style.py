#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将源文档的样式修改为目标文档的样式
用法: python apply_style.py <源文件> <目标样式文件> [输出文件]
"""

import os
import sys
import copy
import argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def copy_run_font(source_run, target_run):
    """复制 run 的字体格式"""
    if target_run.font.name:
        source_run.font.name = target_run.font.name
    try:
        source_run._element.rPr.rFonts.set(qn('w:eastAsia'), target_run.font.name or '宋体')
    except Exception:
        pass
    if target_run.font.size:
        source_run.font.size = target_run.font.size
    if target_run.font.bold is not None:
        source_run.font.bold = target_run.font.bold
    if target_run.font.italic is not None:
        source_run.font.italic = target_run.font.italic
    if target_run.font.underline:
        source_run.font.underline = target_run.font.underline
    if target_run.font.color and target_run.font.color.rgb:
        source_run.font.color.rgb = target_run.font.color.rgb


def copy_paragraph_format(source_para, target_para):
    """复制段落的格式"""
    pf = source_para.paragraph_format
    tpf = target_para.paragraph_format

    if tpf.alignment is not None:
        pf.alignment = tpf.alignment
    if tpf.left_indent is not None:
        pf.left_indent = tpf.left_indent
    if tpf.right_indent is not None:
        pf.right_indent = tpf.right_indent
    if tpf.first_line_indent is not None:
        pf.first_line_indent = tpf.first_line_indent
    if tpf.space_before is not None:
        pf.space_before = tpf.space_before
    if tpf.space_after is not None:
        pf.space_after = tpf.space_after
    if tpf.line_spacing is not None:
        pf.line_spacing = tpf.line_spacing
    if tpf.line_spacing_rule is not None:
        pf.line_spacing_rule = tpf.line_spacing_rule


def copy_table_cell_shading(source_cell, target_cell):
    """复制单元格背景色"""
    def get_shading(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        return shd

    try:
        source_tc = source_cell._tc
        target_tc = target_cell._tc
        source_tcPr = source_tc.get_or_add_tcPr()
        target_tcPr = target_cell._tc.get_or_add_tcPr()
        # 复制背景色
        target_shd = target_tcPr.find(qn('w:shd'))
        if target_shd is not None:
            source_shd = copy.deepcopy(target_shd)
            # 移除旧的 shd
            old = source_tcPr.find(qn('w:shd'))
            if old is not None:
                source_tcPr.remove(old)
            source_tcPr.append(source_shd)
    except Exception:
        pass


def apply_styles(source_doc, target_doc):
    """将目标文档的样式应用到源文档"""
    total_para = len(source_doc.paragraphs)
    target_para_count = len(target_doc.paragraphs)
    total_table = len(source_doc.tables)
    target_table_count = len(target_doc.tables)

    # 遍历源文档所有段落，用目标文档对应段落的样式覆盖
    for i, src_para in enumerate(source_doc.paragraphs):
        if i < target_para_count:
            tgt_para = target_doc.paragraphs[i]
            # 复制段落格式
            copy_paragraph_format(src_para, tgt_para)
            # 复制 run 字体格式（只处理有内容的 run）
            src_runs = [r for r in src_para.runs if r.text.strip()]
            tgt_runs = [r for r in tgt_para.runs if r.text.strip()]
            count = min(len(src_runs), len(tgt_runs))
            for j in range(count):
                copy_run_font(src_runs[j], tgt_runs[j])
        else:
            # 源文档段落多于目标文档时，跳过该段落
            pass

    # 处理表格
    for i, src_table in enumerate(source_doc.tables):
        if i < target_table_count:
            tgt_table = target_doc.tables[i]
            rows = min(len(src_table.rows), len(tgt_table.rows))
            for r in range(rows):
                src_row = src_table.rows[r]
                tgt_row = tgt_table.rows[r]
                cells = min(len(src_row.cells), len(tgt_row.cells))
                for c in range(cells):
                    # 复制单元格格式
                    copy_table_cell_shading(src_row.cells[c], tgt_row.cells[c])
                    # 复制单元格内段落格式
                    src_cell_paragraphs = src_row.cells[c].paragraphs
                    tgt_cell_paragraphs = tgt_row.cells[c].paragraphs
                    count = min(len(src_cell_paragraphs), len(tgt_cell_paragraphs))
                    for p in range(count):
                        copy_paragraph_format(src_cell_paragraphs[p], tgt_cell_paragraphs[p])
                        src_runs = [r for r in src_cell_paragraphs[p].runs if r.text.strip()]
                        tgt_runs = [r for r in tgt_cell_paragraphs[p].runs if r.text.strip()]
                        run_count = min(len(src_runs), len(tgt_runs))
                        for r in range(run_count):
                            copy_run_font(src_runs[r], tgt_runs[r])

    # 处理页眉页脚
    for src_section, tgt_section in zip(source_doc.sections, target_doc.sections):
        for src_header, tgt_header in zip(
            [src_section.header, src_section.first_page_header, src_section.even_page_header],
            [tgt_section.header, tgt_section.first_page_header, tgt_section.even_page_header]
        ):
            if src_header and tgt_header:
                for src_para, tgt_para in zip(src_header.paragraphs, tgt_header.paragraphs):
                    copy_paragraph_format(src_para, tgt_para)
                    src_runs = [r for r in src_para.runs if r.text.strip()]
                    tgt_runs = [r for r in tgt_para.runs if r.text.strip()]
                    count = min(len(src_runs), len(tgt_runs))
                    for j in range(count):
                        copy_run_font(src_runs[j], tgt_runs[j])

        for src_footer, tgt_footer in zip(
            [src_section.footer, src_section.first_page_footer, src_section.even_page_footer],
            [tgt_section.footer, tgt_section.first_page_footer, tgt_section.even_page_footer]
        ):
            if src_footer and tgt_footer:
                for src_para, tgt_para in zip(src_footer.paragraphs, tgt_footer.paragraphs):
                    copy_paragraph_format(src_para, tgt_para)
                    src_runs = [r for r in src_para.runs if r.text.strip()]
                    tgt_runs = [r for r in tgt_para.runs if r.text.strip()]
                    count = min(len(src_runs), len(tgt_runs))
                    for j in range(count):
                        copy_run_font(src_runs[j], tgt_runs[j])


def main():
    parser = argparse.ArgumentParser(description='将目标文档的样式应用到源文档')
    parser.add_argument('source', help='源文件（将被修改）')
    parser.add_argument('target', help='目标样式文件')
    parser.add_argument('-o', '--output', help='输出文件路径（不指定则覆盖源文件）', default=None)
    args = parser.parse_args()

    if not os.path.exists(args.source):
        print(f'错误: 源文件不存在: {args.source}')
        sys.exit(1)
    if not os.path.exists(args.target):
        print(f'错误: 目标样式文件不存在: {args.target}')
        sys.exit(1)

    output_path = args.output or args.source

    try:
        source_doc = Document(args.source)
        target_doc = Document(args.target)
        apply_styles(source_doc, target_doc)
        source_doc.save(output_path)
        print(f'完成！样式已应用，文件保存至: {output_path}')
    except Exception as e:
        print(f'处理失败: {e}')
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()