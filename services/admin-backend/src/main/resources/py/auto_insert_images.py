#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
自动向 Word 文档插入图片

用法: python auto_insert_images.py <文件目录> <图片库目录> <数量n>
功能：
1. 遍历指定目录及子目录下的所有 .docx 文件
2. 从图片库随机选取 n 张图片插入到每个文件中
3. 第一张图片插入在文件最前面居中显示
4. 后面 n-1 张图片随机插入到文末区域的换行段后，居中显示
"""

import os
import sys
import random
import shutil
import argparse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def is_valid_image(filename):
    """判断是否为有效图片文件"""
    valid_exts = ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')
    lower = filename.lower()
    return any(lower.endswith(ext) for ext in valid_exts)


def collect_files(directory, extensions):
    """递归收集目录下的指定类型文件"""
    result = []
    for root, dirs, files in os.walk(directory):
        for f in files:
            if any(f.lower().endswith(ext) for ext in extensions):
                result.append(os.path.join(root, f))
    return result


def add_image_to_document(doc, image_path, centered=True):
    """
    向 doc 添加一张图片，返回是否成功
    - centered=True: 居中对齐
    """
    try:
        # 计算图片宽度（最大不超过 6 英寸，保持比例）
        from PIL import Image as PILImage
        with PILImage.open(image_path) as img:
            orig_w, orig_h = img.size
        max_width = 6.0
        width_inches = min(orig_w / 96.0, max_width)  # 假设 96 DPI

        # 添加图片
        run = doc.add_picture(image_path, width=Inches(width_inches))

        # 获取刚添加的段落（最后一个段落）
        paragraph = doc.paragraphs[-1]
        if centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置图片下方的段前距
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        return True
    except Exception as e:
        print(f"  添加图片失败: {e}")
        return False


def insert_images_to_docx(docx_path, image_pool, count):
    """
    向单个 docx 文件插入图片
    - image_pool: 可用图片路径列表
    - count: 需要插入的图片数量
    返回是否成功
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"  打开文档失败: {e}")
        return False

    if not image_pool:
        print(f"  图片库为空，跳过")
        return False

    # 随机选 count 张图片（不放回）
    selected = []
    pool_copy = list(image_pool)
    for _ in range(min(count, len(pool_copy))):
        if not pool_copy:
            break
        img = random.choice(pool_copy)
        pool_copy.remove(img)
        selected.append(img)

    if len(selected) == 0:
        print(f"  未找到可用图片，跳过")
        return False

    # 第一张图片：插入到最前面，居中
    first_img = selected[0]
    try:
        # 创建居中的图片段落，插入到最开始位置
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn as qname

        # 读取图片尺寸
        from PIL import Image as PILImage
        with PILImage.open(first_img) as img:
            orig_w, orig_h = img.size
        max_width = 6.0
        width_inches = min(orig_w / 96.0, max_width)

        # 在最前面插入一个空的居中段落，然后在里面添加图片
        first_para = doc.paragraphs[0]
        # 在第一个段落之前插入新段落
        new_para = doc.paragraphs[0].insert_paragraph_before()
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_para.add_run()
        run.add_picture(first_img, width=Inches(width_inches))
        new_para.paragraph_format.space_before = Pt(6)
        new_para.paragraph_format.space_after = Pt(6)
    except Exception as e:
        print(f"  插入第一张图片失败: {e}")
        # 尝试普通方式添加
        try:
            first_para = doc.paragraphs[0]
            first_para.insert_paragraph_before()
            doc.add_picture(first_img, width=Inches(6))
        except Exception:
            pass

    # 后面 count-1 张图片：随机插入到文末的换行段后
    remaining = selected[1:]
    if remaining:
        # 找到所有段落
        paragraphs = doc.paragraphs
        # 取后 half 的段落（更容易插入到文末区域）
        start_idx = max(0, len(paragraphs) // 2)
        end_idx = len(paragraphs)
        target_indices = list(range(start_idx, end_idx))

        if target_indices:
            for img_path in remaining:
                try:
                    target_pos = random.choice(target_indices)
                    target_para = paragraphs[target_pos]

                    from PIL import Image as PILImage
                    with PILImage.open(img_path) as img:
                        orig_w, orig_h = img.size
                    max_width = 6.0
                    width_inches = min(orig_w / 96.0, max_width)

                    # 在目标段落后插入新段落
                    new_para = target_para.insert_paragraph_after()
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    new_run = new_para.add_run()
                    new_run.add_picture(img_path, width=Inches(width_inches))
                    new_para.paragraph_format.space_before = Pt(6)
                    new_para.paragraph_format.space_after = Pt(6)

                    # 刷新段落引用
                    paragraphs = doc.paragraphs
                except Exception as e:
                    print(f"  插入后续图片失败: {e}")

    try:
        doc.save(docx_path)
        return True
    except Exception as e:
        print(f"  保存文档失败: {e}")
        return False


def process_directory(file_dir, image_lib_dir, count):
    """处理整个目录"""
    if not os.path.isdir(file_dir):
        print(f"错误: 文件目录不存在: {file_dir}")
        return 0, 0

    if not os.path.isdir(image_lib_dir):
        print(f"错误: 图片库目录不存在: {image_lib_dir}")
        return 0, 0

    # 收集图片
    image_pool = [os.path.join(image_lib_dir, f) for f in os.listdir(image_lib_dir)
                  if is_valid_image(f)]

    # 如果图片库是目录，也递归收集子目录图片
    sub_dirs = [os.path.join(image_lib_dir, d) for d in os.listdir(image_lib_dir)
                if os.path.isdir(os.path.join(image_lib_dir, d))]
    for sub in sub_dirs:
        sub_images = [os.path.join(sub, f) for f in os.listdir(sub) if is_valid_image(f)]
        image_pool.extend(sub_images)

    image_pool = list(set(image_pool))  # 去重
    print(f"找到 {len(image_pool)} 张可用图片")

    if len(image_pool) == 0:
        print("错误: 图片库中没有找到有效图片（支持 jpg/jpeg/png/gif/bmp/webp）")
        return 0, 0

    # 收集 docx 文件
    docx_files = collect_files(file_dir, ('.docx',))
    doc_files = collect_files(file_dir, ('.doc',))

    if not docx_files and not doc_files:
        print("错误: 文件目录中未找到 .doc/.docx 文件")
        return 0, 0

    print(f"找到 {len(docx_files)} 个 .docx 文件，{len(doc_files)} 个 .doc 文件（将跳过 .doc）")

    success = 0
    failed = 0
    for f in docx_files:
        rel = os.path.relpath(f, file_dir)
        print(f"处理: {rel} ...", end=' ')
        if insert_images_to_docx(f, image_pool, count):
            print("完成")
            success += 1
        else:
            print("失败")
            failed += 1

    return success, failed


def main():
    parser = argparse.ArgumentParser(description='自动向 Word 文档插入图片')
    parser.add_argument('fileDir', help='文件目录（doc/docx 文件所在目录）')
    parser.add_argument('imageLibDir', help='图片库目录')
    parser.add_argument('count', type=int, help='每个文件插入的图片数量')
    args = parser.parse_args()

    if args.count < 1:
        print("错误: 图片数量必须 >= 1")
        sys.exit(1)

    success, failed = process_directory(args.fileDir, args.imageLibDir, args.count)
    print(f"\n处理完成！成功: {success}，失败: {failed}")


if __name__ == '__main__':
    main()