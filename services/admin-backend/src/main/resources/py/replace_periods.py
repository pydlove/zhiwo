#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量替换 Word 文档中的句号、分号、顿号、冒号和中文引号

功能：
1. 遍历指定目录下的所有 .docx 文件
2. 将句号（。）、分号（；）、顿号（、）、中文冒号（：）、英文冒号（:）
   替换为逗号（，）
3. 将中文引号「」替换为空
4. 将破折号（——）替换为逗号（，）
5. 例外：如果标点后面紧跟换行符（\n）或位于段落末尾，则保留
6. 支持自定义规则（--custom-rules JSON数组）
7. 孤立短句合并：逗号<=2且句号==1的段落 → 句号改逗号并合并到下一段（跳过中间空段落）
"""

import os
import sys
import argparse
import json
import re
from docx import Document


def replace_periods(text, custom_rules=None):
    """
    将文本中的句号（。）、分号（；）、顿号（、）、中文冒号（：）、英文冒号（:）
    替换为逗号（，），将中文引号「」替换为空，将破折号（——）替换为逗号（，）。
    保留后面紧跟换行符（\n）或位于文本末尾的标点。
    custom_rules: 可选，自定义规则列表 [[from, to], [from, to], ...]
    """
    result = []
    i = 0
    while i < len(text):
        # 优先匹配破折号 ——（两个中文破折号）
        if i + 1 < len(text) and text[i] == '—' and text[i + 1] == '—':
            if i + 2 >= len(text) or text[i + 2] == '\n':
                result.append('——')
            else:
                result.append('，')
            i += 2
        elif text[i] == '「' or text[i] == '」':
            result.append('')
            i += 1
        elif text[i] in ('。', '；', '：', ':'):
            if i + 1 >= len(text) or text[i + 1] == '\n':
                result.append(text[i])
            else:
                result.append('，')
            i += 1
        elif text[i] == '、':
            if i + 1 >= len(text) or text[i + 1] == '\n':
                result.append(text[i])
            else:
                result.append('，')
            i += 1
        else:
            result.append(text[i])
            i += 1
    text = ''.join(result)

    # 应用自定义规则
    if custom_rules:
        for rule in custom_rules:
            if isinstance(rule, list) and len(rule) >= 2:
                from_str = rule[0]
                to_str = rule[1] if rule[1] else ''
                text = text.replace(from_str, to_str)

    return text


def _delete_paragraph(paragraph):
    """从文档中删除段落（底层 XML 操作）"""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _apply_paragraph_text(paragraph, new_text):
    """将处理后的文本安全地应用到段落（尽量保留格式）"""
    if not new_text:
        paragraph.clear()
        return
    # 如果段落只有一个 run，直接修改
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = new_text
        return
    # 多个 runs：清空后重建
    paragraph.clear()
    paragraph.add_run(new_text)


def _remove_empty_paragraphs(paragraphs):
    """删除段落列表中的所有空段落（从后往前删，避免索引错乱）"""
    to_remove = []
    for i, p in enumerate(paragraphs):
        if not p.text.strip():
            to_remove.append(i)
    for idx in reversed(to_remove):
        _delete_paragraph(paragraphs[idx])


def _process_paragraphs_with_context(paragraphs, custom_rules):
    """跨段落处理孤立短句：
    对于逗号<=2且句号==1的段落，将句号改为逗号，并跳过中间的空段落合并到下一段。
    """
    if not paragraphs:
        return

    # 收集所有段落文本
    texts = [p.text for p in paragraphs]

    # 计算合并：逗号 <= 2 且句号 == 1 的段落，句号改逗号后合并到下一段
    i = 0
    merge_indices = set()  # 记录需要删除的段落索引
    while i < len(texts) - 1:
        segment = texts[i]
        comma_count = segment.count('，')
        period_count = segment.count('。')

        if comma_count <= 2 and period_count == 1:
            # 句号变逗号
            segment = segment.replace('。', '，', 1)

            # 跳过中间的空段落，找到下一个有内容的段落
            j = i + 1
            while j < len(texts) and not texts[j].strip():
                merge_indices.add(j)
                j += 1

            if j < len(texts):
                # 合并到下一个有内容的段落
                texts[j] = segment + texts[j]

            # 标记当前段待删除
            merge_indices.add(i)
            # 继续从下一个非空段落检查（合并后的新段可能也满足条件）
            i = j
        else:
            i += 1

    # 从后往前删除被合并的段落
    for idx in sorted(merge_indices, reverse=True):
        _delete_paragraph(paragraphs[idx])

    # 应用 replace_periods 并写回（跳过已删除的）
    for i, p in enumerate(paragraphs):
        if i in merge_indices:
            continue
        new_text = replace_periods(texts[i], custom_rules)
        _apply_paragraph_text(p, new_text)


def process_document(input_path, output_path, custom_rules=None):
    """处理单个文档"""
    try:
        doc = Document(input_path)

        # 跨段落处理孤立短句
        _process_paragraphs_with_context(list(doc.paragraphs), custom_rules)

        # 处理表格内的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _process_paragraphs_with_context(list(cell.paragraphs), custom_rules)

        # 处理页眉页脚
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    _process_paragraphs_with_context(list(header.paragraphs), custom_rules)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    _process_paragraphs_with_context(list(footer.paragraphs), custom_rules)

        # 去除所有空段落（全局清理空行）
        _remove_empty_paragraphs(list(doc.paragraphs))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_paras = list(cell.paragraphs)
                    _remove_empty_paragraphs(cell_paras)
                    # 确保单元格至少保留一个空段落，否则 Word 结构异常
                    if not list(cell.paragraphs):
                        cell.add_paragraph()
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    _remove_empty_paragraphs(list(header.paragraphs))
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    _remove_empty_paragraphs(list(footer.paragraphs))

        doc.save(output_path)
        return True

    except Exception as e:
        print(f"  错误: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(
        description='批量替换 Word 文档中的句号（保留换行前的句号）'
    )
    parser.add_argument('directory', help='目标目录路径')
    parser.add_argument(
        '-o', '--output',
        help='输出目录（不指定则覆盖原文件）',
        default=None
    )
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='试运行，只显示会处理的文件而不实际修改'
    )
    parser.add_argument(
        '--custom-rules',
        help='自定义规则JSON数组，格式: [["from","to"], ["from","to"]]',
        default=None
    )

    args = parser.parse_args()

    # 解析自定义规则
    custom_rules = None
    if args.custom_rules:
        try:
            custom_rules = json.loads(args.custom_rules)
            print(f"自定义规则: {custom_rules}")
        except Exception as e:
            print(f"警告: 自定义规则JSON解析失败，将使用内置规则: {e}")

    target_dir = os.path.abspath(args.directory)
    if not os.path.isdir(target_dir):
        print(f"错误: '{target_dir}' 不是有效的目录")
        sys.exit(1)

    output_dir = None
    if args.output:
        output_dir = os.path.abspath(args.output)
        os.makedirs(output_dir, exist_ok=True)

    # 收集文件
    docx_files = []
    doc_files = []

    for root, dirs, files in os.walk(target_dir):
        for file in files:
            file_path = os.path.join(root, file)
            if file.lower().endswith('.docx'):
                docx_files.append(file_path)
            elif file.lower().endswith('.doc'):
                doc_files.append(file_path)

    # 处理 .docx 文件
    processed = 0
    failed = 0
    skipped = 0

    print(f"\n找到 {len(docx_files)} 个 .docx 文件")
    if doc_files:
        print(f"注意: 跳过 {len(doc_files)} 个 .doc 文件（请先转换为 .docx）")

    for file_path in docx_files:
        rel_path = os.path.relpath(file_path, target_dir)

        if output_dir:
            # 保持目录结构
            rel_dir = os.path.dirname(rel_path)
            out_subdir = os.path.join(output_dir, rel_dir)
            os.makedirs(out_subdir, exist_ok=True)
            out_path = os.path.join(output_dir, rel_path)
        else:
            out_path = file_path

        if args.dry_run:
            print(f"[试运行] 将处理: {rel_path}")
            processed += 1
            continue

        print(f"处理中: {rel_path} ...", end=' ')
        if process_document(file_path, out_path, custom_rules):
            print("完成")
            processed += 1
        else:
            failed += 1

    print(f"\n{'='*50}")
    print(f"处理完成!")
    print(f"  成功: {processed}")
    if failed:
        print(f"  失败: {failed}")
    if doc_files:
        print(f"  跳过 .doc: {len(doc_files)} (请转换为 .docx 后处理)")
    print(f"{'='*50}")


if __name__ == '__main__':
    main()
